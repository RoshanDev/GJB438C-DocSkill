from __future__ import annotations

from pathlib import Path
import textwrap

ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "skills" / "gjb438c-md-first" / "gjb438c_suite"
PKG.mkdir(parents=True, exist_ok=True)


def write(relative: str, content: str) -> None:
    path = ROOT / relative
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(textwrap.dedent(content).lstrip(), encoding="utf-8")


write(
    "skills/gjb438c-md-first/gjb438c_suite/profile_quality.py",
    r'''
    from __future__ import annotations

    from copy import deepcopy
    from dataclasses import asdict, dataclass, field, is_dataclass
    from functools import lru_cache
    import json
    from pathlib import Path
    import re
    from typing import Any, Iterable

    import yaml

    from .markdown_doc import MarkdownDocument, parse_markdown
    from .quality import audit_markdown

    PROFILE_DIR = Path(__file__).resolve().parent / "data" / "profiles"
    TIER_ALIASES = {
        "prototype": "prototype",
        "proto": "prototype",
        "small": "standard",
        "standard": "standard",
        "medium": "large",
        "large": "large",
        "very-large": "critical",
        "very_large": "critical",
        "critical": "critical",
    }
    VALID_TIERS = ("prototype", "standard", "large", "critical")


    class ProfileQualityError(RuntimeError):
        pass


    @dataclass(frozen=True, slots=True)
    class ProfileIssue:
        severity: str
        code: str
        message: str
        line: int | None = None
        artifact_kind: str | None = None
        artifact_id: str | None = None

        def as_dict(self) -> dict[str, Any]:
            return asdict(self)


    @dataclass(slots=True)
    class ProfileAuditReport:
        document_type: str
        tier: str
        audit_profile: str
        issues: list[ProfileIssue] = field(default_factory=list)
        counts: dict[str, int] = field(default_factory=dict)
        minimums: dict[str, int] = field(default_factory=dict)
        heading_coverage_percent: float = 0.0

        @property
        def errors(self) -> list[ProfileIssue]:
            return [item for item in self.issues if item.severity == "ERROR"]

        @property
        def passed(self) -> bool:
            return not self.errors

        def as_dict(self) -> dict[str, Any]:
            return {
                "document_type": self.document_type,
                "tier": self.tier,
                "audit_profile": self.audit_profile,
                "passed": self.passed,
                "summary": {
                    "errors": len(self.errors),
                    "warnings": sum(1 for item in self.issues if item.severity == "WARN"),
                    "heading_coverage_percent": round(self.heading_coverage_percent, 2),
                },
                "counts": self.counts,
                "minimums": self.minimums,
                "issues": [item.as_dict() for item in self.issues],
            }

        def to_json(self) -> str:
            return json.dumps(self.as_dict(), ensure_ascii=False, indent=2)

        def to_text(self) -> str:
            state = "PASS" if self.passed else "FAIL"
            lines = [
                f"[{state}] profile {self.document_type} tier={self.tier} "
                f"headings={self.heading_coverage_percent:.2f}% errors={len(self.errors)}"
            ]
            for item in self.issues:
                where = ""
                if item.artifact_kind:
                    where += f" kind={item.artifact_kind}"
                if item.artifact_id:
                    where += f" id={item.artifact_id}"
                if item.line:
                    where += f" line={item.line}"
                lines.append(f"- {item.severity} {item.code}:{where} {item.message}")
            return "\n".join(lines)


    @dataclass(slots=True)
    class CombinedAuditReport:
        generic: Any
        profile: ProfileAuditReport

        @property
        def passed(self) -> bool:
            return bool(getattr(self.generic, "passed", False)) and self.profile.passed

        def as_dict(self) -> dict[str, Any]:
            try:
                generic = json.loads(self.generic.to_json())
            except Exception:
                generic = {
                    "passed": bool(getattr(self.generic, "passed", False)),
                    "text": self.generic.to_text(),
                }
            return {
                "passed": self.passed,
                "generic": generic,
                "profile": self.profile.as_dict(),
            }

        def to_json(self) -> str:
            return json.dumps(self.as_dict(), ensure_ascii=False, indent=2)

        def to_text(self) -> str:
            return f"{self.generic.to_text()}\n{self.profile.to_text()}"


    def normalize_tier(value: str | None) -> str:
        key = str(value or "prototype").strip().lower().replace(" ", "-")
        try:
            return TIER_ALIASES[key]
        except KeyError as exc:
            raise ProfileQualityError(
                f"tier/scale 必须是 {', '.join(VALID_TIERS)}；实际为 {value!r}"
            ) from exc


    @lru_cache(maxsize=None)
    def _load_profile_mapping_cached(code: str) -> dict[str, Any]:
        normalized = str(code).strip().upper()
        path = PROFILE_DIR / f"{normalized.lower()}.yaml"
        if not path.is_file():
            raise ProfileQualityError(f"找不到内置 Profile：{normalized} ({path})")
        value = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
        if not isinstance(value, dict):
            raise ProfileQualityError(f"Profile 必须是 YAML 映射：{path}")
        if str(value.get("code", "")).upper() != normalized:
            raise ProfileQualityError(f"Profile code 不一致：{path}")
        return value


    def load_profile_mapping(code: str) -> dict[str, Any]:
        """每次返回新的深拷贝，调用方不能污染后续审计。"""
        return deepcopy(_load_profile_mapping_cached(str(code).strip().upper()))


    def tier_minimum(value: Any, tier: str, *, default: int = 0) -> int:
        if isinstance(value, bool):
            return int(value)
        if isinstance(value, (int, float)):
            return max(0, int(value))
        if isinstance(value, dict):
            normalized = normalize_tier(tier)
            if normalized in value:
                return max(0, int(value[normalized]))
            for alias, canonical in TIER_ALIASES.items():
                if canonical == normalized and alias in value:
                    return max(0, int(value[alias]))
            raise ProfileQualityError(
                f"Profile minimum 缺少 tier={normalized}；可用键={sorted(value)}"
            )
        if value is None:
            return max(0, int(default))
        raise ProfileQualityError(f"无效 minimum：{value!r}")


    def resolve_document_tier(document: MarkdownDocument, explicit: str | None = None) -> str:
        if explicit:
            return normalize_tier(explicit)
        for container_name in ("quality", "project"):
            container = document.metadata.get(container_name)
            if not isinstance(container, dict):
                continue
            for key in ("tier", "scale"):
                if container.get(key):
                    return normalize_tier(str(container[key]))
        return "prototype"


    def has_declared_tier(document: MarkdownDocument, explicit: str | None = None) -> bool:
        if explicit:
            return True
        for container_name in ("quality", "project"):
            container = document.metadata.get(container_name)
            if isinstance(container, dict) and (container.get("tier") or container.get("scale")):
                return True
        return False


    def is_fixture(document: MarkdownDocument) -> bool:
        quality = document.metadata.get("quality")
        return isinstance(quality, dict) and quality.get("fixture") is True


    def document_type_from_metadata(document: MarkdownDocument) -> str | None:
        value = document.metadata.get("document")
        if isinstance(value, dict) and value.get("type"):
            return str(value["type"]).strip().upper()
        for key in ("document_type", "type"):
            if document.metadata.get(key):
                return str(document.metadata[key]).strip().upper()
        return None


    def artifact_mapping(artifact: Any) -> dict[str, Any]:
        for name in ("data", "payload", "value", "content"):
            value = getattr(artifact, name, None)
            if isinstance(value, dict):
                return value
        if isinstance(artifact, dict):
            return artifact
        if is_dataclass(artifact):
            value = asdict(artifact)
            for name in ("data", "payload", "value", "content"):
                nested = value.get(name)
                if isinstance(nested, dict):
                    return nested
            return value
        return {}


    def artifact_kind(artifact: Any) -> str:
        value = getattr(artifact, "kind", None)
        if value is None and isinstance(artifact, dict):
            value = artifact.get("kind")
        return str(value or "").strip().lower().removeprefix("gjb-")


    def artifact_line(artifact: Any) -> int | None:
        value = getattr(artifact, "line", None)
        if value is None:
            value = getattr(artifact, "start_line", None)
        try:
            return int(value) if value is not None else None
        except (TypeError, ValueError):
            return None


    def _blank(value: Any) -> bool:
        if value is None:
            return True
        if isinstance(value, str):
            return not value.strip()
        if isinstance(value, (list, tuple, set, dict)):
            return not value
        return False


    def _normalized_heading(value: str) -> str:
        value = re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", value.strip())
        return re.sub(r"[\s\u3000:：()（）]+", "", value).lower()


    def _dynamic_heading(value: str) -> bool:
        compact = value.upper()
        return bool(re.search(r"(?:^|[.\s])X(?:[.\s]|$)|(?:^|[.\s])Y(?:[.\s]|$)", compact))


    def _iter_artifacts(document: MarkdownDocument) -> Iterable[Any]:
        value = getattr(document, "artifacts", ())
        return value if isinstance(value, (list, tuple)) else tuple(value or ())


    def audit_profile_document(
        source: str | Path | MarkdownDocument,
        *,
        document_type: str | None = None,
        audit_profile: str = "review",
        tier: str | None = None,
    ) -> ProfileAuditReport:
        document = source if isinstance(source, MarkdownDocument) else parse_markdown(source)
        detected = document_type_from_metadata(document)
        code = str(document_type or detected or "").strip().upper()
        if not code:
            raise ProfileQualityError("无法确定文档类型；需设置 document.type 或 --type")
        mapping = load_profile_mapping(code)
        resolved_tier = resolve_document_tier(document, tier)
        report = ProfileAuditReport(code, resolved_tier, audit_profile)
        severity = "ERROR" if audit_profile == "release" else "WARN"

        if detected and detected != code:
            report.issues.append(
                ProfileIssue("ERROR", "PROFILE_TYPE_MISMATCH", f"文档声明 {detected}，实际按 {code} 审计")
            )
        if audit_profile == "release" and not has_declared_tier(document, tier):
            report.issues.append(
                ProfileIssue("ERROR", "PROFILE_TIER_REQUIRED", "发布文档必须显式声明 quality.tier/scale")
            )
        if audit_profile == "release" and resolved_tier == "prototype" and not is_fixture(document):
            report.issues.append(
                ProfileIssue(
                    "ERROR",
                    "PROFILE_PROTOTYPE_RELEASE_FORBIDDEN",
                    "prototype 仅供示例和 CI；生产发布必须使用 standard、large 或 critical",
                )
            )
        allowed = mapping.get("release_rules", {}).get("allowed_tiers", VALID_TIERS)
        allowed = {normalize_tier(str(item)) for item in allowed}
        if audit_profile == "release" and resolved_tier not in allowed:
            report.issues.append(
                ProfileIssue(
                    "ERROR",
                    "PROFILE_TIER_NOT_ALLOWED",
                    f"{code} 不允许以 tier={resolved_tier} 发布；允许值={sorted(allowed)}",
                )
            )

        artifacts = list(_iter_artifacts(document))
        by_kind: dict[str, list[Any]] = {}
        seen_ids: dict[str, tuple[str, int | None]] = {}
        for artifact in artifacts:
            kind = artifact_kind(artifact)
            if kind:
                by_kind.setdefault(kind, []).append(artifact)
            payload = artifact_mapping(artifact)
            artifact_id = str(payload.get("id", "")).strip()
            if artifact_id:
                if artifact_id in seen_ids:
                    prior_kind, prior_line = seen_ids[artifact_id]
                    report.issues.append(
                        ProfileIssue(
                            "ERROR",
                            "PROFILE_DUPLICATE_ID",
                            f"稳定标识重复；首次 kind={prior_kind} line={prior_line}",
                            artifact_line(artifact),
                            kind,
                            artifact_id,
                        )
                    )
                else:
                    seen_ids[artifact_id] = (kind, artifact_line(artifact))

        contracts = mapping.get("artifact_contracts", [])
        if not isinstance(contracts, list):
            raise ProfileQualityError(f"{code} artifact_contracts 必须是列表")
        for contract in contracts:
            if not isinstance(contract, dict) or not contract.get("kind"):
                continue
            kind = str(contract["kind"]).strip().lower().removeprefix("gjb-")
            items = by_kind.get(kind, [])
            minimum = tier_minimum(contract.get("minimum"), resolved_tier)
            report.counts[kind] = len(items)
            report.minimums[kind] = minimum
            if len(items) < minimum:
                report.issues.append(
                    ProfileIssue(
                        severity,
                        "PROFILE_ARTIFACT_COUNT_LOW",
                        f"gjb-{kind} 数量 {len(items)} 低于 {code}/{resolved_tier} 最低值 {minimum}",
                        artifact_kind=kind,
                    )
                )
            required_fields = tuple(str(item) for item in contract.get("required_fields", []) if item)
            for artifact in items:
                payload = artifact_mapping(artifact)
                artifact_id = str(payload.get("id", "")).strip() or None
                for field_name in required_fields:
                    if _blank(payload.get(field_name)):
                        report.issues.append(
                            ProfileIssue(
                                severity,
                                "PROFILE_REQUIRED_FIELD_MISSING",
                                f"gjb-{kind} 缺少字段 {field_name}",
                                artifact_line(artifact),
                                kind,
                                artifact_id,
                            )
                        )

        expected = []
        for item in mapping.get("outline", []):
            if not isinstance(item, dict):
                continue
            title = str(item.get("title", "")).strip()
            if title and not _dynamic_heading(title):
                expected.append((int(item.get("level", 1)), _normalized_heading(title), title))
        actual = {
            _normalized_heading(str(getattr(item, "title", item)))
            for item in getattr(document, "headings", ())
        }
        matched = sum(1 for _, normalized, _ in expected if normalized in actual)
        report.heading_coverage_percent = matched / len(expected) * 100 if expected else 100.0
        for level, normalized, title in expected:
            if normalized in actual:
                continue
            missing_severity = "ERROR" if audit_profile == "release" and level == 1 else severity
            report.issues.append(
                ProfileIssue(
                    missing_severity,
                    "PROFILE_HEADING_MISSING",
                    f"缺少 Profile 章节：{title}",
                )
            )
        return report


    def audit_markdown_with_profile(
        source: str | Path,
        *,
        profile: str = "review",
        document_type: str | None = None,
        baseline_srs: str | Path | None = None,
        tier: str | None = None,
        scale: str | None = None,
    ) -> CombinedAuditReport:
        generic = audit_markdown(
            source,
            profile=profile,
            document_type=document_type,
            baseline_srs=baseline_srs,
        )
        specific = audit_profile_document(
            source,
            document_type=document_type,
            audit_profile=profile,
            tier=tier or scale,
        )
        return CombinedAuditReport(generic, specific)
    ''',
)

write(
    "skills/gjb438c-md-first/gjb438c_suite/volume.py",
    r'''
    from __future__ import annotations

    from collections import Counter
    from dataclasses import asdict, dataclass
    import hashlib
    import json
    import math
    from pathlib import Path
    import re
    import shutil
    import subprocess
    import tempfile
    from typing import Any

    from pypdf import PdfReader

    from .markdown_doc import MarkdownDocument, parse_markdown, strip_quality_blocks
    from .profile_quality import (
        ProfileQualityError,
        is_fixture,
        load_profile_mapping,
        normalize_tier,
        resolve_document_tier,
    )


    class VolumeError(RuntimeError):
        pass


    @dataclass(frozen=True, slots=True)
    class RenderedPageMetrics:
        total_pages: int
        body_start_page: int
        body_pages: int
        visible_characters: int
        thin_pages: int
        thin_page_ratio: float
        duplicate_pages: int
        duplicate_page_ratio: float
        minimum_page_characters: int

        def as_dict(self) -> dict[str, Any]:
            return asdict(self)


    @dataclass(frozen=True, slots=True)
    class VolumeResult:
        document_type: str
        tier: str
        total_pages: int
        body_start_page: int
        body_pages: int
        minimum_body_pages: int
        visible_characters: int
        minimum_visible_characters: int
        effective_units: int
        duplicate_prose_ratio: float
        duplicate_page_ratio: float
        thin_page_ratio: float
        passed: bool
        issues: tuple[str, ...]
        source_sha256: str
        docx_sha256: str

        def as_dict(self) -> dict[str, Any]:
            return asdict(self)

        def to_json(self) -> str:
            return json.dumps(self.as_dict(), ensure_ascii=False, indent=2)

        def to_text(self) -> str:
            state = "PASS" if self.passed else "FAIL"
            lines = [
                f"[{state}] volume {self.document_type} tier={self.tier} "
                f"body_pages={self.body_pages}/{self.minimum_body_pages} "
                f"visible_chars={self.visible_characters}/{self.minimum_visible_characters} "
                f"thin={self.thin_page_ratio:.2%} duplicate_pages={self.duplicate_page_ratio:.2%} "
                f"duplicate_prose={self.duplicate_prose_ratio:.2%}"
            ]
            lines.extend(f"- {item}" for item in self.issues)
            return "\n".join(lines)


    def sha256_file(path: str | Path) -> str:
        digest = hashlib.sha256()
        with Path(path).open("rb") as stream:
            for block in iter(lambda: stream.read(1024 * 1024), b""):
                digest.update(block)
        return digest.hexdigest()


    def sha256_text(text: str) -> str:
        return hashlib.sha256(text.encode("utf-8")).hexdigest()


    def resolve_tier(document: MarkdownDocument, explicit: str | None = None) -> str:
        return resolve_document_tier(document, explicit)


    def _policy(document_type: str, tier: str) -> dict[str, Any]:
        mapping = load_profile_mapping(document_type)
        try:
            value = mapping["volume_policy"][normalize_tier(tier)]
        except KeyError as exc:
            raise VolumeError(f"缺少 {document_type}/{tier} 体量策略") from exc
        if not isinstance(value, dict):
            raise VolumeError(f"无效体量策略：{document_type}/{tier}")
        return value


    def volume_policy(document_type: str, tier: str) -> dict[str, Any]:
        normalized = normalize_tier(tier)
        result = dict(_policy(document_type.upper(), normalized))
        result.update({"document_type": document_type.upper(), "tier": normalized})
        return result


    def minimum_body_pages(
        document_type: str,
        tier: str,
        override: int | None = None,
    ) -> int:
        normalized = normalize_tier(tier)
        floor = int(_policy(document_type.upper(), normalized).get("minimum_pages", 0))
        if override is None:
            return floor
        requested = int(override)
        if requested < floor:
            raise VolumeError(
                f"--min-body-pages 只能提高下限；{document_type.upper()}/{normalized} "
                f"策略下限为 {floor}，不能降为 {requested}"
            )
        return requested


    def _visible_markdown_text(document: MarkdownDocument) -> str:
        body = strip_quality_blocks(document.body)
        body = re.sub(r"```.*?```", "", body, flags=re.S)
        body = re.sub(r"!\[[^]]*\]\([^)]*\)", "", body)
        body = re.sub(r"<[^>]+>", "", body)
        return body


    def visible_markdown_characters(document: MarkdownDocument) -> int:
        return len(re.sub(r"\s+", "", _visible_markdown_text(document)))


    def _normalized_prose_paragraphs(document: MarkdownDocument) -> list[str]:
        result: list[str] = []
        for block in re.split(r"\n\s*\n+", _visible_markdown_text(document)):
            lines: list[str] = []
            for line in block.splitlines():
                stripped = line.strip()
                if not stripped or stripped.startswith("#"):
                    continue
                if re.match(r"^\|(?:\s*:?-{3,}:?\s*\|)+\s*$", stripped):
                    continue
                lines.append(stripped)
            normalized = re.sub(r"[\s\W_]+", "", "".join(lines), flags=re.UNICODE)
            if len(normalized) >= 80:
                result.append(normalized)
        return result


    def duplicate_prose_metrics(document: MarkdownDocument) -> tuple[int, float, int]:
        paragraphs = _normalized_prose_paragraphs(document)
        if not paragraphs:
            return 0, 0.0, 0
        counts = Counter(paragraphs)
        total = sum(len(item) for item in paragraphs)
        duplicate = sum((count - 1) * len(item) for item, count in counts.items())
        return duplicate, (duplicate / total if total else 0.0), max(counts.values(), default=0)


    def _table_count(document: MarkdownDocument) -> int:
        return sum(
            1
            for line in _visible_markdown_text(document).splitlines()
            if re.match(r"^\s*\|(?:\s*:?-{3,}:?\s*\|)+\s*$", line)
        )


    def _figure_count(document: MarkdownDocument) -> int:
        return len(re.findall(r"!\[[^]]*\]\([^)]*\)", strip_quality_blocks(document.body)))


    def effective_units(document: MarkdownDocument) -> tuple[int, int]:
        visible = visible_markdown_characters(document)
        duplicate, _, _ = duplicate_prose_metrics(document)
        deduplicated = max(0, visible - duplicate)
        artifacts = len(tuple(getattr(document, "artifacts", ()) or ()))
        units = deduplicated + _table_count(document) * 350 + _figure_count(document) * 450 + artifacts * 120
        return visible, units


    def _scaled_visible_floor(document_type: str, tier: str, body_pages: int) -> int:
        policy = _policy(document_type, tier)
        base_pages = max(1, int(policy.get("minimum_pages", 1)))
        base_chars = max(0, int(policy.get("minimum_visible_characters", 0)))
        return max(base_chars, math.ceil(base_chars / base_pages * body_pages))


    def markdown_volume_issues(
        document: MarkdownDocument,
        document_type: str,
        tier: str,
        audit_profile: str,
        *,
        min_body_pages_override: int | None = None,
    ) -> list[dict[str, Any]]:
        normalized = normalize_tier(tier)
        policy = _policy(document_type, normalized)
        floor = minimum_body_pages(document_type, normalized, min_body_pages_override)
        visible, units = effective_units(document)
        visible_floor = _scaled_visible_floor(document_type, normalized, floor)
        _, duplicate_ratio, max_repeat = duplicate_prose_metrics(document)
        duplicate_limit = float(policy.get("maximum_duplicate_page_ratio", 0.08))
        severity = "ERROR" if audit_profile == "release" else "WARN"
        issues: list[dict[str, Any]] = []
        quality = document.metadata.get("quality")
        declared = isinstance(quality, dict) and (quality.get("tier") or quality.get("scale"))
        if audit_profile == "release" and not declared:
            issues.append(
                {"severity": "ERROR", "code": "VOLUME_TIER_REQUIRED", "message": "发布文档必须显式声明 quality.tier/scale"}
            )
        if audit_profile == "release" and normalized == "prototype" and not is_fixture(document):
            issues.append(
                {
                    "severity": "ERROR",
                    "code": "VOLUME_PROTOTYPE_RELEASE_FORBIDDEN",
                    "message": "prototype 仅供仓库示例/CI；生产发布必须选择 standard、large 或 critical",
                }
            )
        if visible < visible_floor:
            issues.append(
                {
                    "severity": severity,
                    "code": "VOLUME_VISIBLE_CONTENT_TOO_THIN",
                    "message": f"可见正文字符 {visible} 低于 {document_type}/{normalized} 最低值 {visible_floor}",
                }
            )
        if max_repeat >= 3 and duplicate_ratio > duplicate_limit:
            issues.append(
                {
                    "severity": severity,
                    "code": "VOLUME_DUPLICATE_CONTENT",
                    "message": f"重复长段落占比 {duplicate_ratio:.2%} 超过 {duplicate_limit:.2%}，最大重复次数 {max_repeat}",
                }
            )
        if units < visible_floor:
            issues.append(
                {
                    "severity": severity,
                    "code": "VOLUME_EFFECTIVE_CONTENT_TOO_THIN",
                    "message": f"去重后的等效内容单位 {units} 低于 {visible_floor}",
                }
            )
        return issues


    def _office() -> str:
        candidates = (
            shutil.which("libreoffice"),
            shutil.which("soffice"),
            shutil.which("soffice.exe"),
            r"C:\Program Files\LibreOffice\program\soffice.exe",
        )
        for value in candidates:
            if value and Path(value).exists():
                return str(value)
        raise VolumeError("release 页数门禁需要 LibreOffice/soffice，不能用 DOCX 元数据代替真实渲染")


    def _normalized_page_text(text: str) -> str:
        lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped or re.fullmatch(r"[-—–]?\s*\d+\s*[-—–]?", stripped):
                continue
            lines.append(stripped)
        return re.sub(r"[\s\W_]+", "", "".join(lines), flags=re.UNICODE)


    def _looks_like_toc(text: str) -> bool:
        compact = text.replace(" ", "")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        page_refs = sum(bool(re.search(r"(?:\.{2,}|…{2,}|\s)\d+\s*$", line)) for line in lines)
        return "目录" in compact or page_refs >= 5


    def _estimate_body_start(page_texts: list[str]) -> int:
        if not page_texts:
            return 0
        start = min(3, len(page_texts) - 1)
        for index in range(start, len(page_texts)):
            text = page_texts[index]
            if _looks_like_toc(text):
                continue
            normalized = _normalized_page_text(text)
            if re.search(r"(?:^|[^0-9])1范围", normalized) or normalized.startswith("范围"):
                return index
            if index >= 8 and len(normalized) >= 180:
                return index
        return min(4, len(page_texts) - 1)


    def rendered_page_metrics(
        docx: str | Path,
        *,
        body_start_page: int | None = None,
        minimum_page_characters: int = 100,
    ) -> RenderedPageMetrics:
        source = Path(docx).resolve()
        if not source.is_file():
            raise FileNotFoundError(source)
        with tempfile.TemporaryDirectory(prefix="gjb438c-pages-") as name:
            out = Path(name)
            profile = out / "lo-profile"
            profile.mkdir()
            command = [
                _office(),
                f"-env:UserInstallation={profile.as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out),
                str(source),
            ]
            result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=300)
            pdf = out / f"{source.stem}.pdf"
            if result.returncode or not pdf.is_file():
                raise VolumeError("LibreOffice 页数渲染失败：" + (result.stderr.strip() or result.stdout.strip()))
            reader = PdfReader(str(pdf))
            page_texts = [(page.extract_text() or "") for page in reader.pages]
        total = len(page_texts)
        if body_start_page is None:
            start_index = _estimate_body_start(page_texts)
        else:
            if body_start_page < 1 or body_start_page > max(1, total):
                raise VolumeError(f"body_start_page 超出范围：{body_start_page}/{total}")
            start_index = body_start_page - 1
        body = page_texts[start_index:]
        normalized = [_normalized_page_text(item) for item in body]
        visible = sum(len(item) for item in normalized)
        thin = sum(len(item) < minimum_page_characters for item in normalized)
        eligible = [item for item in normalized if len(item) >= minimum_page_characters]
        counts = Counter(eligible)
        duplicate = sum(count - 1 for count in counts.values())
        body_pages = len(body)
        return RenderedPageMetrics(
            total_pages=total,
            body_start_page=start_index + 1,
            body_pages=body_pages,
            visible_characters=visible,
            thin_pages=thin,
            thin_page_ratio=(thin / body_pages if body_pages else 1.0),
            duplicate_pages=duplicate,
            duplicate_page_ratio=(duplicate / body_pages if body_pages else 1.0),
            minimum_page_characters=minimum_page_characters,
        )


    def audit_rendered_volume(
        source: str | Path | MarkdownDocument,
        document_type: str,
        docx: str | Path,
        *,
        tier: str | None = None,
        min_body_pages_override: int | None = None,
        body_start_page: int | None = None,
    ) -> VolumeResult:
        document = source if isinstance(source, MarkdownDocument) else parse_markdown(source)
        normalized = resolve_tier(document, tier)
        policy = _policy(document_type, normalized)
        floor = minimum_body_pages(document_type, normalized, min_body_pages_override)
        base_pages = max(1, int(policy.get("minimum_pages", 1)))
        base_chars = max(0, int(policy.get("minimum_visible_characters", 0)))
        per_page = max(80, min(240, math.floor(base_chars / base_pages * 0.35)))
        metrics = rendered_page_metrics(
            docx,
            body_start_page=body_start_page,
            minimum_page_characters=per_page,
        )
        visible_floor = _scaled_visible_floor(document_type, normalized, floor)
        _, effective = effective_units(document)
        _, duplicate_prose_ratio, _ = duplicate_prose_metrics(document)
        max_thin = float(policy.get("maximum_thin_page_ratio", 0.2))
        max_duplicate = float(policy.get("maximum_duplicate_page_ratio", 0.08))
        issues: list[str] = []
        if metrics.body_pages < floor:
            issues.append(f"正文页数 {metrics.body_pages} 低于 {document_type}/{normalized} 发布下限 {floor}")
        if metrics.visible_characters < visible_floor:
            issues.append(f"Word 可见正文字符 {metrics.visible_characters} 低于最低值 {visible_floor}")
        if metrics.thin_page_ratio > max_thin:
            issues.append(f"薄页比例 {metrics.thin_page_ratio:.2%} 超过 {max_thin:.2%}")
        if metrics.duplicate_page_ratio > max_duplicate:
            issues.append(f"重复页比例 {metrics.duplicate_page_ratio:.2%} 超过 {max_duplicate:.2%}")
        if duplicate_prose_ratio > max_duplicate:
            issues.append(f"Markdown 重复长段落比例 {duplicate_prose_ratio:.2%} 超过 {max_duplicate:.2%}")
        return VolumeResult(
            document_type=document_type.upper(),
            tier=normalized,
            total_pages=metrics.total_pages,
            body_start_page=metrics.body_start_page,
            body_pages=metrics.body_pages,
            minimum_body_pages=floor,
            visible_characters=metrics.visible_characters,
            minimum_visible_characters=visible_floor,
            effective_units=effective,
            duplicate_prose_ratio=round(duplicate_prose_ratio, 6),
            duplicate_page_ratio=round(metrics.duplicate_page_ratio, 6),
            thin_page_ratio=round(metrics.thin_page_ratio, 6),
            passed=not issues,
            issues=tuple(issues),
            source_sha256=sha256_text(document.raw),
            docx_sha256=sha256_file(docx),
        )
    ''',
)

write(
    "skills/gjb438c-md-first/gjb438c_suite/evidence.py",
    r'''
    from __future__ import annotations

    from collections import defaultdict
    from pathlib import Path
    from typing import Any

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    from .markdown_doc import MarkdownDocument, parse_markdown
    from .profile_quality import artifact_kind, artifact_mapping


    def _text(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "是" if value else "否"
        if isinstance(value, dict):
            return "；".join(f"{key}={_text(item)}" for key, item in value.items())
        if isinstance(value, (list, tuple, set)):
            return "；".join(_text(item) for item in value)
        return str(value)


    def _format_run(run, *, bold: bool = False, size: float = 10.5) -> None:
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体" if bold else "宋体")


    def _format_paragraph(paragraph, *, bold: bool = False, size: float = 10.5) -> None:
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            _format_run(run, bold=bold, size=size)


    def append_evidence_appendix(
        docx: str | Path,
        source: str | Path | MarkdownDocument,
    ) -> Path:
        document = source if isinstance(source, MarkdownDocument) else parse_markdown(source)
        artifacts = list(getattr(document, "artifacts", ()) or ())
        if not artifacts:
            return Path(docx)
        target = Path(docx)
        word = Document(str(target))
        word.add_page_break()
        heading = word.add_paragraph("结构化工程证据（机器可审计视图）")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _format_paragraph(heading, bold=True, size=12)
        intro = word.add_paragraph(
            "本附录将 Markdown 中的 gjb-* 证据块转换为 Word 可见内容。"
            "它用于评审、追踪和机器审计，不替代正文中的技术论证。"
        )
        _format_paragraph(intro)
        groups: dict[str, list[Any]] = defaultdict(list)
        for artifact in artifacts:
            groups[artifact_kind(artifact) or "unknown"].append(artifact)
        for kind in sorted(groups):
            sub = word.add_paragraph(f"gjb-{kind}")
            _format_paragraph(sub, bold=True, size=10.5)
            table = word.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            headers = ("序号", "稳定标识", "字段", "内容")
            for cell, value in zip(table.rows[0].cells, headers):
                cell.text = value
                _format_paragraph(cell.paragraphs[0], bold=True)
            row_number = 0
            for artifact in groups[kind]:
                payload = artifact_mapping(artifact)
                artifact_id = str(payload.get("id", ""))
                fields = [(key, value) for key, value in payload.items() if key != "id"] or [("内容", "")]
                for field_name, value in fields:
                    row_number += 1
                    cells = table.add_row().cells
                    cells[0].text = str(row_number)
                    cells[1].text = artifact_id
                    cells[2].text = str(field_name)
                    cells[3].text = _text(value)
                    for cell in cells:
                        _format_paragraph(cell.paragraphs[0])
        word.save(str(target))
        return target
    ''',
)

print("runtime core A written")
