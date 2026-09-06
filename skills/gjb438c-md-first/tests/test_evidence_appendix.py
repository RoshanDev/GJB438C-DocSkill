from pathlib import Path

from docx import Document

from gjb438c_suite.evidence import append_evidence_appendix
from gjb438c_suite.markdown_doc import parse_markdown
from gjb438c_suite.profile_quality import artifact_mapping


def test_structured_evidence_is_visible_in_word(tmp_path: Path):
    source = Path(__file__).resolve().parents[1] / "examples/SRS.example.md"
    parsed = parse_markdown(source)
    artifacts = list(parsed.artifacts)
    assert artifacts
    first_id = str(artifact_mapping(artifacts[0]).get("id", ""))
    assert first_id
    docx = tmp_path / "evidence.docx"
    Document().save(docx)
    append_evidence_appendix(docx, parsed)
    word = Document(docx)
    visible = "\n".join(
        [paragraph.text for paragraph in word.paragraphs]
        + [cell.text for table in word.tables for row in table.rows for cell in row.cells]
    )
    assert "结构化工程证据" in visible
    assert first_id in visible
