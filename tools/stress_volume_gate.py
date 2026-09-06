"""Synthetic pagination/duplicate-detector stress test, NOT a quality sample."""
from pathlib import Path
import tempfile
from docx import Document
from gjb438c_suite.volume import rendered_page_metrics


def build(path, repeated):
    doc=Document()
    for title in ('Cover','Signatures','Changes','Table of contents'):
        doc.add_paragraph(title);doc.add_page_break()
    for i in range(300):
        value = 0 if repeated else i
        doc.add_paragraph((f'Synthetic pagination fixture {value:04d}; this is not project evidence. ')*8)
        if i!=299:doc.add_page_break()
    doc.save(path)


def main():
    with tempfile.TemporaryDirectory(prefix='gjb-stress-') as name:
        root=Path(name)
        for repeated in (False,True):
            path=root/f'{repeated}.docx';build(path,repeated)
            metrics=rendered_page_metrics(path,body_start_page=5)
            print(metrics.as_dict())
            assert metrics.body_pages==300
            if repeated:assert metrics.duplicate_page_ratio>0.95
            else:assert metrics.duplicate_page_ratio==0
    return 0


if __name__=='__main__':raise SystemExit(main())
