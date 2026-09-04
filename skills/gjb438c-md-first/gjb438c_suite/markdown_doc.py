from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Any, Iterable

import yaml
from docx import Document
from docx.oxml.ns import qn

FRONT_MATTER_RE = re.compile(r"\A---\s*\n(?P<yaml>.*?)\n---\s*\n", re.DOTALL)
HEADING_RE = re.compile(r"^(?P<marks>#{1,9})\s+(?P<title>.+?)\s*$", re.MULTILINE)
FENCE_RE = re.compile(
    r"^```(?P<lang>gjb-[a-z0-9_-]+)\s*\n(?P<body>.*?)^```\s*$",
    re.MULTILINE | re.DOTALL | re.IGNORECASE,
)
ANY_FENCE_RE = re.compile(
    r"^```(?P<lang>[^\n]*)\n(?P<body>.*?)^```\s*$",
    re.MULTILINE | re.DOTALL,
)
PLACEHOLDER_RE = re.compile(
    r"(?:\bTODO\b|\bTBD\b|待补充|待确认|待定|XXXX+|"
    r"[<\[]\s*(?:TBD|TODO|待补充|待确认)\s*[>\]]|\{\{[^}\n]+\}\})",
    re.IGNORECASE,
)
NUMBER_PREFIX_RE = re.compile(r"^\s*(?:第\s*)?\d+(?:\.\d+)*(?:\s*[章节条、.]|\s+)")
TOC_PAGE_SUFFIX_RE = re.compile(r"(?:\t|\s{2,})\d+\s*$")


@dataclass(slots=True)
class Heading:
    level: int
    title: str
    line: int
    number: str | None = None


@dataclass(slots=True)
class Artifact:
    kind: str
    data: dict[str, Any]
    line: int
    language: str

    @property
    def identifier(self) -> str | None:
        value = self.data.get("id")
        return str(value).strip() if value is not None else None


@dataclass(slots=True)
class Placeholder:
    text: str
    line: int


@dataclass(slots=True)
class MarkdownDocument:
    path: Path
    raw: str
    metadata: dict[str, Any]
    body: str
    headings: list[Heading] = field(default_factory=list)
    artifacts: list[Artifact] = field(default_factory=list)
    placeholders: list[Placeholder] = field(default_factory=list)
    parse_errors: list[str] = field(default_factory=list)

    def artifacts_of(self, kind: str) -> list[Artifact]:
        normalized = kind.lower().replace("_", "-")
        return [artifact for artifact in self.artifacts if artifact.kind == normalized]

    @property
    def visible_body(self) -> str:
        return strip_quality_blocks(self.body)


def line_number(text: str, offset: int) -> int:
    return text.count("\n", 0, offset) + 1


def nested_get(data: dict[str, Any], dotted: str, default: Any = None) -> Any:
    current: Any = data
    for part in dotted.split("."):
        if not isinstance(current, dict) or part not in current:
            return default
        current = current[part]
    return current


def split_front_matter(raw: str) -> tuple[dict[str, Any], str, int, list[str]]:
    metadata: dict[str, Any] = {}
    errors: list[str] = []
    body = raw
    offset = 0
    match = FRONT_MATTER_RE.match(raw)
    if not match:
        return metadata, body, offset, ["缺少 YAML front matter（文件开头应为 ---）"]
    try:
        loaded = yaml.safe_load(match.group("yaml")) or {}
        if not isinstance(loaded, dict):
            errors.append("YAML front matter 必须是映射对象")
        else:
            metadata = loaded
    except yaml.YAMLError as exc:
        errors.append(f"YAML front matter 解析失败：{exc}")
    offset = match.end()
    body = raw[offset:]
    return metadata, body, offset, errors


def parse_markdown(path: str | Path) -> MarkdownDocument:
    source_path = Path(path)
    raw = source_path.read_text(encoding="utf-8")
    metadata, body, body_offset, errors = split_front_matter(raw)

    headings = [
        Heading(
            len(match.group("marks")),
            match.group("title").strip(),
            line_number(raw, body_offset + match.start()),
            None,
        )
        for match in HEADING_RE.finditer(body)
    ]

    artifacts: list[Artifact] = []
    for match in FENCE_RE.finditer(body):
        language = match.group("lang").lower()
        kind = language.removeprefix("gjb-").replace("_", "-")
        source_line = line_number(raw, body_offset + match.start())
        try:
            data = yaml.safe_load(match.group("body")) or {}
            if not isinstance(data, dict):
                errors.append(f"第 {source_line} 行 {language} 数据块必须是 YAML 映射")
                continue
            explicit_kind = data.get("kind")
            if explicit_kind:
                kind = str(explicit_kind).lower().replace("_", "-")
            artifacts.append(Artifact(kind, data, source_line, language))
        except yaml.YAMLError as exc:
            errors.append(f"第 {source_line} 行 {language} YAML 解析失败：{exc}")

    placeholders = [
        Placeholder(match.group(0), line_number(raw, match.start()))
        for match in PLACEHOLDER_RE.finditer(raw)
    ]
    return MarkdownDocument(
        source_path, raw, metadata, body, headings, artifacts, placeholders, errors
    )


def strip_quality_blocks(body: str) -> str:
    visible = FENCE_RE.sub("", body)
    # Quality blocks are the Markdown source-of-truth for automated review, not
    # delivery prose. When a terminal appendix contains only those blocks, drop
    # its now-empty heading from the rendered Word document instead of shipping
    # a visually blank appendix page/section.
    visible = re.sub(
        r"(?ms)^#{1,9}\s+附录\s*[A-Za-zＡ-Ｚａ-ｚ0-9一二三四五六七八九十]*\s*质量门禁数据块\s*\Z",
        "",
        visible,
    )
    return visible.strip() + "\n"


def strip_number_prefix(title: str) -> str:
    return NUMBER_PREFIX_RE.sub("", title).strip()


def _style_outline_map(document: Document) -> dict[str, int]:
    """Map style IDs to OOXML outline levels (1-based)."""
    result: dict[str, int] = {}
    styles = document.styles.element
    for style in styles.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId"))
        outline = style.find("./w:pPr/w:outlineLvl", {"w": qn("w:p").split("}")[0][1:]})
        if style_id and outline is not None:
            value = outline.get(qn("w:val"))
            if value is not None:
                result[style_id] = int(value) + 1
    return result


def _toc_level(style_name: str) -> int | None:
    match = re.search(r"(?:toc|目录)\s*([1-9])", style_name, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _heading_level(style_name: str, style_id: str | None, outline: dict[str, int]) -> int | None:
    match = re.search(r"(?:heading|标题)\s*([1-9])", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    if style_id and style_id in outline:
        return outline[style_id]
    return None


def extract_template_outline(template_path: str | Path) -> list[Heading]:
    """Extract a document outline, preferring the real Word TOC block.

    The 20 repository templates are not stylistically identical. Their TOC is
    the most stable source of hierarchy because it already contains the
    template's intended chapter and clause sequence. A body-heading fallback is
    used when no TOC entries are present.
    """
    document = Document(Path(template_path))
    outline_map = _style_outline_map(document)

    toc: list[Heading] = []
    collecting = False
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        level = _toc_level(style_name)
        if level is not None and value:
            collecting = True
            value = TOC_PAGE_SUFFIX_RE.sub("", value).strip()
            if value.replace(" ", "") == "目录":
                continue
            number_match = re.match(r"^(\d+(?:\.\d+)*)\s+(.+)$", value)
            number = number_match.group(1) if number_match else None
            title = number_match.group(2).strip() if number_match else value
            toc.append(Heading(level, title, 0, number))
        elif collecting and toc and value and level is None:
            # First visible non-TOC paragraph marks the end of the field result.
            break

    candidates = toc
    if not candidates:
        candidates = []
        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if not value:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            style_id = paragraph.style.style_id if paragraph.style is not None else None
            level = _heading_level(style_name, style_id, outline_map)
            typed = re.match(r"^(\d+(?:\.\d+)*)\s*[、.．]?\s*(.+)$", value)
            number = typed.group(1) if typed else None
            title = typed.group(2).strip() if typed else strip_number_prefix(value)
            if level is None and number:
                level = number.count(".") + 1
            if level is None or not title:
                continue
            candidates.append(Heading(min(level, 9), title, 0, number))

    # Remove front matter and duplicate TOC/body entries while preserving order.
    start = next(
        (
            index
            for index, heading in enumerate(candidates)
            if heading.level == 1 and re.search(r"范围|引言|概述", heading.title)
        ),
        0,
    )
    candidates = candidates[start:]
    result: list[Heading] = []
    seen: set[tuple[int, str, str | None]] = set()
    for heading in candidates:
        title = heading.title.strip()
        if title.replace(" ", "") in {"目录", "签字页", "变更履历", "修改页"}:
            continue
        key = (heading.level, title, heading.number)
        if key in seen:
            continue
        seen.add(key)
        result.append(Heading(heading.level, title, 0, heading.number))
    if not result:
        raise ValueError(f"模板中未识别到正文标题：{template_path}")
    return result


def _artifact_stub(kind: str, index: int) -> str:
    identifier = f"{kind.upper().replace('-', '_')}-{index:03d}"
    specialized: dict[str, list[str]] = {
        "requirement": [
            f"id: REQ-{index:03d}",
            "statement: 待补充（使用‘应’陈述）",
            "rationale: 待补充",
            "source: 待补充",
            "priority: P1",
            "verification: [待补充]",
            "acceptance: [待补充]",
        ],
        "design-unit": [
            f"id: DU-{index:03d}",
            "requirements: [待补充]",
            "responsibility: 待补充",
            "behavior: 待补充",
            "interfaces: [待补充]",
            "data: [待补充]",
            "states: [待补充]",
            "errors: 待补充",
            "concurrency: 待补充",
            "security: 待补充",
            "deployment: 待补充",
            "verification: [待补充]",
            "source_refs: [待补充]",
        ],
        "decision": [
            f"id: ADR-{index:03d}",
            "context: 待补充",
            "options: [待补充]",
            "decision: 待补充",
            "rationale: 待补充",
            "consequences: 待补充",
            "status: proposed",
            "source_refs: [待补充]",
        ],
        "architecture": [
            f"id: ARCH-{index:03d}",
            "components: [待补充]",
            "connectors: [待补充]",
            "deployment: 待补充",
            "failure_domains: 待补充",
            "source_refs: [待补充]",
        ],
        "interface": [
            f"id: IF-{index:03d}",
            "provider: 待补充",
            "consumer: 待补充",
            "protocol: 待补充",
            "input: 待补充",
            "output: 待补充",
            "timing: 待补充",
            "errors: 待补充",
            "security: 待补充",
            "compatibility: 待补充",
            "source_refs: [待补充]",
        ],
        "data": [
            f"id: DM-{index:03d}",
            "owner: 待补充",
            "schema: 待补充",
            "constraints: 待补充",
            "transaction: 待补充",
            "retention: 待补充",
            "security: 待补充",
            "recovery: 待补充",
            "source_refs: [待补充]",
        ],
        "scenario": [
            f"id: SCN-{index:03d}",
            "requirements: [待补充]",
            "trigger: 待补充",
            "preconditions: 待补充",
            "steps: [待补充]",
            "failures: [待补充]",
            "postconditions: 待补充",
            "observability: 待补充",
            "source_refs: [待补充]",
        ],
        "deployment": [
            f"id: DEPLOY-{index:03d}",
            "nodes: 待补充",
            "placement: 待补充",
            "resources: 待补充",
            "network: 待补充",
            "storage: 待补充",
            "upgrade: 待补充",
            "rollback: 待补充",
            "source_refs: [待补充]",
        ],
        "security": [
            f"id: SEC-{index:03d}",
            "assets: [待补充]",
            "threats: [待补充]",
            "controls: [待补充]",
            "audit: 待补充",
            "residual_risk: 待补充",
            "source_refs: [待补充]",
        ],
        "verification": [
            f"id: VT-{index:03d}",
            "target: 待补充",
            "method: 待补充",
            "criteria: 待补充",
            "evidence: 待补充",
            "source_refs: [待补充]",
        ],
        "traceability": [
            f"id: TRACE-{index:03d}",
            "source_refs: [待补充]",
            "requirements: [待补充]",
            "forward_targets: [待补充]",
        ],
    }
    rows = specialized.get(
        kind,
        [
            f"id: {identifier}",
            "title: 待补充",
            "source_refs: [待补充]",
            "content: 待补充",
        ],
    )
    return f"```gjb-{kind}\n" + "\n".join(rows) + "\n```\n"


def render_skeleton(
    *,
    document_type: Any,
    outline: Iterable[Heading],
    project: dict[str, Any] | None = None,
) -> str:
    project = project or {}
    software = project.get("software", {}) if isinstance(project.get("software"), dict) else {}
    document = project.get("document", {}) if isinstance(project.get("document"), dict) else {}
    front = project.get("front_matter", {}) if isinstance(project.get("front_matter"), dict) else {}
    metadata = {
        "document": {
            "type": document_type.code,
            "title": document.get("title", f"{software.get('name', '待补充软件')}{document_type.chinese_name}"),
            "id": document.get("id", f"DOC-{document_type.code}-TBD"),
            "version": document.get("version", "V0.1"),
            "status": document.get("status", "draft"),
            "standard_clause": document_type.clause,
            "appendix": document_type.appendix,
        },
        "software": {
            "name": software.get("name", "待补充"),
            "version": software.get("version", "待补充"),
            "identifier": software.get("identifier", "待补充"),
        },
        "organization": project.get("organization", "待补充"),
        "classification": project.get("classification", "待补充"),
        "date": project.get("date", "待补充"),
        "front_matter": {
            "template": front.get("template", "templates/front-matter/standard-front-matter.docx"),
            "archive_id": front.get("archive_id", "待补充"),
            "project_code": front.get("project_code", "待补充"),
            "phase": front.get("phase", "待补充"),
            "date_cn": front.get("date_cn", "待补充"),
        },
        "signatures": project.get(
            "signatures",
            {
                "prepared": {"name": "", "date": ""},
                "reviewed": {"name": "", "date": ""},
                "standard_reviewed": {"name": "", "date": ""},
                "countersigned": {"name": "", "date": ""},
                "approved": {"name": "", "date": ""},
            },
        ),
        "revisions": project.get(
            "revisions",
            [{"date": "待补充", "version": "V0.1", "description": "建立文档骨架", "author": "待补充"}],
        ),
        "sources": project.get("sources", []),
    }
    yaml_text = yaml.safe_dump(metadata, allow_unicode=True, sort_keys=False).rstrip()
    lines = ["---", yaml_text, "---", "", f"# {metadata['document']['title']}", ""]
    counters = [0] * 9
    for heading in outline:
        level = min(max(heading.level, 1), 9)
        number = heading.number
        if number is None and not re.match(r"^(附录|参考文献|索引)", heading.title):
            counters[level - 1] += 1
            for index in range(level, len(counters)):
                counters[index] = 0
            number = ".".join(str(value) for value in counters[:level] if value)
        marks = "#" * min(level, 6)
        label = f"{number} {heading.title}" if number else heading.title
        lines.extend([f"{marks} {label}", "", "待补充。", ""])
    lines.extend(["# 附录A 质量门禁数据块", ""])
    for index, kind in enumerate(document_type.required_artifacts, 1):
        lines.append(_artifact_stub(kind, index))
    return "\n".join(lines).rstrip() + "\n"
