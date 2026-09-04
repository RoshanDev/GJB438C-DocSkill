from __future__ import annotations

from pathlib import Path
import shutil

import pytest
from docx import Document

from gjb438c_suite.audit_docx import audit_docx
from gjb438c_suite.finalize import FinalizeError, refresh_toc_cache
from gjb438c_suite.import_word import import_word
from gjb438c_suite.render import render_document


def test_render_audit_and_exact_markdown_round_trip(tmp_path: Path) -> None:
    root = Path(__file__).resolve().parents[1]
    source = root / "examples/SRS.example.md"
    output = tmp_path / "SRS.docx"
    render_document(source, output, profile="release")

    review = audit_docx(output, profile="review")
    assert review.passed, review.to_text()
    assert any(issue.code == "TOC_CACHE_STALE" for issue in review.warnings)

    release = audit_docx(output, profile="release")
    assert not release.passed
    assert any(issue.code == "TOC_CACHE_STALE" for issue in release.errors)

    returned = tmp_path / "SRS.returned.md"
    result = import_word(output, returned)
    assert result.exact_round_trip
    assert returned.read_text(encoding="utf-8") == source.read_text(encoding="utf-8")


def test_word_body_change_produces_review_candidate(tmp_path: Path) -> None:
    root = Path(__file__).resolve().parents[1]
    source = root / "examples/SRS.example.md"
    output = tmp_path / "SRS.docx"
    render_document(source, output, profile="release")

    document = Document(output)
    paragraph = next(p for p in document.paragraphs if "长时任务" in p.text)
    paragraph.add_run("（Word 评审修改）")
    document.save(output)

    returned = tmp_path / "candidate.md"
    result = import_word(output, returned)
    assert not result.exact_round_trip
    assert result.warning
    text = returned.read_text(encoding="utf-8")
    assert "requires_review: true" in text
    assert "Word 评审修改" in text


@pytest.mark.skipif(shutil.which("libreoffice") is None and shutil.which("soffice") is None, reason="LibreOffice unavailable")
def test_refresh_toc_transplants_cache_without_rewriting_document(tmp_path: Path) -> None:
    root = Path(__file__).resolve().parents[1]
    source = root / "examples/SRS.example.md"
    output = tmp_path / "SRS.docx"
    render_document(source, output, profile="release")
    try:
        refresh_toc_cache(output)
    except FinalizeError as exc:
        pytest.skip(str(exc))

    report = audit_docx(output, profile="release")
    assert report.passed, report.to_text()
    assert not report.metrics["toc_cache_stale"]

    returned = tmp_path / "SRS.returned.md"
    result = import_word(output, returned)
    assert result.exact_round_trip
    assert returned.read_text(encoding="utf-8") == source.read_text(encoding="utf-8")
