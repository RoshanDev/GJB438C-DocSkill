from __future__ import annotations

from pathlib import Path
import re
import textwrap

ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "skills/gjb438c-md-first"
PKG = SKILL / "gjb438c_suite"
TESTS = SKILL / "tests"

volume_py = r'''from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
import hashlib
import json
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any

import yaml
from pypdf import PdfReader

from .markdown_doc import MarkdownDocument, strip_quality_blocks
from .profiles import DocumentProfile

POLICY_PATH = Path(__file__).resolve().parent / "data" / "volume-policy.yaml"
VALID_SCALES = ("prototype", "small", "medium", "large", "very-large")


class VolumeError(RuntimeError):
    pass


@dataclass(frozen=True, slots=True)
class VolumeResult:
    document_type: str
    scale: str
    pages: int
    minimum_pages: int
    visible_chars: int
    effective_units: int
    minimum_effective_units: int
    passed: bool
    issues: tuple[str, ...]
    policy_version: int
    source_sha256: str
    docx_sha256: str

    def as_dict(self) -> dict[str, Any]:
        return {
            "document_type": self.document_type,
            "scale": self.scale,
            "pages": self.pages,
            "minimum_pages": self.minimum_pages,
            "visible_chars": self.visible_chars,
            "effective_units": self.effective_units,
            "minimum_effective_units": self.minimum_effective_units,
            "passed": self.passed,
            "issues": list(self.issues),
            "policy_version": self.policy_version,
            "source_sha256": self.source_sha256,
            "docx_sha256": self.docx_sha256,
        }

    def to_json(self) -> str:
        return json.dumps(self.as_dict(), ensure_ascii=False, indent=2)

    def to_text(self) -> str:
        state = "PASS" if self.passed else "FAIL"
        tail = "\n- " + "\n- ".join(self.issues) if self.issues else ""
        return (
            f"[{state}] volume {self.document_type} scale={self.scale} "
            f"pages={self.pages}/{self.minimum_pages} "
            f"effective_units={self.effective_units}/{self.minimum_effective_units}{tail}"
        )


def sha256_file(path: str | Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def load_volume_policy() -> dict[str, Any]:
    value = yaml.safe_load(POLICY_PATH.read_text(encoding="utf-8")) or {}
    if not isinstance(value, dict):
        raise VolumeError(f"无效体量策略：{POLICY_PATH}")
    return value


def resolve_scale(document: MarkdownDocument, explicit: str | None = None) -> str:
    value = explicit
    if not value:
        quality = document.metadata.get("quality")
        project = document.metadata.get("project")
        if isinstance(quality, dict):
            value = quality.get("scale")
        if not value and isinstance(project, dict):
            value = project.get("scale")
    normalized = str(value or "prototype").strip().lower().replace("_", "-")
    if normalized not in VALID_SCALES:
        raise VolumeError(f"scale 必须是 {', '.join(VALID_SCALES)}，实际为 {normalized!r}")
    return normalized


def has_declared_scale(document: MarkdownDocument, explicit: str | None = None) -> bool:
    if explicit:
        return True
    for key in ("quality", "project"):
        value = document.metadata.get(key)
        if isinstance(value, dict) and value.get("scale"):
            return True
    return False


def minimum_pages(document_type: str, scale: str, override: int | None = None) -> int:
    if override is not None:
        return max(0, int(override))
    try:
        return int(load_volume_policy()["documents"][document_type.upper()][scale])
    except KeyError as exc:
        raise VolumeError(f"缺少 {document_type}/{scale} 页数策略") from exc


def _visible_chars(document: MarkdownDocument) -> int:
    body = strip_quality_blocks(document.body)
    body = re.sub(r"```.*?```", "", body, flags=re.S)
    body = re.sub(r"!\[[^]]*\]\([^)]*\)", "", body)
    return len(re.sub(r"\s+", "", body))


def _table_count(document: MarkdownDocument) -> int:
    lines = strip_quality_blocks(document.body).splitlines()
    return sum(
        1
        for line in lines
        if re.match(r"^\s*\|(?:\s*:?-{3,}:?\s*\|)+\s*$", line)
    )


def _figure_count(document: MarkdownDocument) -> int:
    return len(re.findall(r"!\[[^]]*\]\([^)]*\)", strip_quality_blocks(document.body)))


def effective_units(document: MarkdownDocument) -> tuple[int, int]:
    chars = _visible_chars(document)
    units = (
        chars
        + _table_count(document) * 700
        + _figure_count(document) * 550
        + len(document.artifacts) * 180
    )
    return chars, units


def markdown_volume_issues(
    document: MarkdownDocument,
    profile: DocumentProfile,
    scale: str,
    audit_profile: str,
    scale_declared: bool,
) -> list[dict[str, Any]]:
    policy = load_volume_policy()
    severity = "ERROR" if audit_profile == "release" else "WARN"
    issues: list[dict[str, Any]] = []
    if audit_profile == "release" and not scale_declared:
        issues.append(
            {
                "severity": "ERROR",
                "code": "VOLUME_SCALE_REQUIRED",
                "message": "发布文档必须在 quality.scale/project.scale 或 --scale 中声明项目规模",
            }
        )
    pages = minimum_pages(profile.document_type.code, scale)
    _, units = effective_units(document)
    floor = pages * int(policy.get("text_equivalent_chars_per_page", 260))
    if units < floor:
        issues.append(
            {
                "severity": severity,
                "code": "VOLUME_EVIDENCE_TOO_THIN",
                "message": (
                    f"{profile.document_type.code}/{scale} 等效内容单位 {units} 低于 {floor}；"
                    "页数不能用空白或重复内容替代"
                ),
            }
        )
    expected = {h.title.strip() for h in profile.outline if h.title.strip()}
    actual = {h.title.strip() for h in document.headings if h.title.strip()}
    coverage = len(expected & actual) / len(expected) * 100 if expected else 100
    required_coverage = float(policy["minimum_heading_coverage_percent"][scale])
    if coverage + 1e-9 < required_coverage:
        issues.append(
            {
                "severity": severity,
                "code": "VOLUME_OUTLINE_COVERAGE_LOW",
                "message": f"Profile 章节覆盖率 {coverage:.2f}% 低于 {required_coverage:.2f}%",
            }
        )
    counts = Counter(a.kind for a in document.artifacts)
    repeatable = policy.get("repeatable_artifact_minima", {})
    for kind in profile.artifact_contracts:
        if kind in repeatable:
            required = int(repeatable[kind][scale])
            if counts[kind] < required:
                issues.append(
                    {
                        "severity": severity,
                        "code": "VOLUME_ARTIFACT_COUNT_LOW",
                        "message": (
                            f"gjb-{kind} 数量 {counts[kind]} 低于 "
                            f"{profile.document_type.code}/{scale} 最低值 {required}"
                        ),
                    }
                )
    return issues


def _office() -> str:
    candidates = (
        shutil.which("libreoffice"),
        shutil.which("soffice"),
        shutil.which("soffice.exe"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
    )
    for value in candidates:
        if value and Path(value).exists():
            return str(value)
    raise VolumeError("页数门禁需要 LibreOffice/soffice；release 不允许跳过真实渲染页数")


def rendered_page_count(docx: str | Path) -> int:
    source = Path(docx).resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    with tempfile.TemporaryDirectory(prefix="gjb438c-pages-") as name:
        out = Path(name)
        result = subprocess.run(
            [_office(), "--headless", "--convert-to", "pdf", "--outdir", str(out), str(source)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=180,
        )
        pdf = out / f"{source.stem}.pdf"
        if result.returncode or not pdf.is_file():
            raise VolumeError(
                "LibreOffice 页数渲染失败："
                + (result.stderr.strip() or result.stdout.strip())
            )
        return len(PdfReader(str(pdf)).pages)


def audit_rendered_volume(
    document: MarkdownDocument,
    profile: DocumentProfile,
    docx: str | Path,
    *,
    scale: str,
    min_pages_override: int | None = None,
) -> VolumeResult:
    docx_path = Path(docx)
    pages = rendered_page_count(docx_path)
    floor = minimum_pages(profile.document_type.code, scale, min_pages_override)
    chars, units = effective_units(document)
    policy = load_volume_policy()
    unit_floor = floor * int(policy.get("text_equivalent_chars_per_page", 260))
    issues: list[str] = []
    if pages < floor:
        issues.append(
            f"渲染页数 {pages} 低于 {profile.document_type.code}/{scale} 发布下限 {floor}"
        )
    if units < unit_floor:
        issues.append(f"等效内容单位 {units} 低于反灌水下限 {unit_floor}")
    return VolumeResult(
        profile.document_type.code,
        scale,
        pages,
        floor,
        chars,
        units,
        unit_floor,
        not issues,
        tuple(issues),
        int(policy.get("policy_version", 1)),
        sha256_text(document.raw),
        sha256_file(docx_path),
    )
'''
(PKG / "volume.py").write_text(volume_py, encoding="utf-8")

suite_py = r'''from __future__ import annotations

from dataclasses import dataclass, field
import json
from pathlib import Path
from typing import Any

import yaml

from .audit_docx import audit_docx
from .markdown_doc import nested_get, parse_markdown
from .profile_quality import audit_markdown_with_profile
from .registry import get_document_type, iter_document_types
from .volume import (
    VALID_SCALES,
    load_volume_policy,
    minimum_pages,
    sha256_file,
    sha256_text,
)


class SuiteError(RuntimeError):
    pass


@dataclass(slots=True)
class SuiteIssue:
    severity: str
    code: str
    message: str
    document_type: str | None = None

    def as_dict(self) -> dict[str, Any]:
        return {
            "severity": self.severity,
            "code": self.code,
            "message": self.message,
            "document_type": self.document_type,
        }


@dataclass(slots=True)
class SuiteDocumentResult:
    document_type: str
    markdown: str
    docx: str
    volume: str
    pages: int = 0
    minimum_pages: int = 0
    passed: bool = False

    def as_dict(self) -> dict[str, Any]:
        return {
            "document_type": self.document_type,
            "markdown": self.markdown,
            "docx": self.docx,
            "volume": self.volume,
            "pages": self.pages,
            "minimum_pages": self.minimum_pages,
            "passed": self.passed,
        }


@dataclass(slots=True)
class SuiteAuditReport:
    manifest: Path
    scale: str
    issues: list[SuiteIssue] = field(default_factory=list)
    documents: list[SuiteDocumentResult] = field(default_factory=list)
    total_pages: int = 0
    minimum_total_pages: int = 0

    @property
    def errors(self) -> list[SuiteIssue]:
        return [issue for issue in self.issues if issue.severity == "ERROR"]

    @property
    def passed(self) -> bool:
        return not self.errors and all(item.passed for item in self.documents)

    def as_dict(self) -> dict[str, Any]:
        return {
            "manifest": str(self.manifest),
            "scale": self.scale,
            "passed": self.passed,
            "summary": {
                "errors": len(self.errors),
                "documents": len(self.documents),
                "total_pages": self.total_pages,
                "minimum_total_pages": self.minimum_total_pages,
            },
            "documents": [item.as_dict() for item in self.documents],
            "issues": [issue.as_dict() for issue in self.issues],
        }

    def to_json(self) -> str:
        return json.dumps(self.as_dict(), ensure_ascii=False, indent=2)

    def to_text(self) -> str:
        state = "PASS" if self.passed else "FAIL"
        lines = [
            f"[{state}] suite {self.manifest} scale={self.scale}",
            (
                f"documents={len(self.documents)} pages="
                f"{self.total_pages}/{self.minimum_total_pages} errors={len(self.errors)}"
            ),
        ]
        for item in self.documents:
            mark = "PASS" if item.passed else "FAIL"
            lines.append(
                f"- {mark} {item.document_type}: pages={item.pages}/{item.minimum_pages}"
            )
        for issue in self.issues:
            prefix = f" {issue.document_type}" if issue.document_type else ""
            lines.append(f"- {issue.severity} {issue.code}:{prefix} {issue.message}")
        return "\n".join(lines)


def _add(
    report: SuiteAuditReport,
    severity: str,
    code: str,
    message: str,
    document_type: str | None = None,
) -> None:
    report.issues.append(SuiteIssue(severity, code, message, document_type))


def _mapping(value: Any, label: str) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise SuiteError(f"{label} 必须是 YAML 映射")
    return value


def _path(base: Path, value: Any, label: str) -> Path:
    if not isinstance(value, str) or not value.strip():
        raise SuiteError(f"{label} 缺少路径")
    path = Path(value.strip())
    return path if path.is_absolute() else (base / path).resolve()


def _volume_payload(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise SuiteError(f"无法读取体量报告 {path}: {exc}") from exc
    return _mapping(value, f"体量报告 {path}")


def audit_suite_manifest(
    manifest: str | Path,
    *,
    scale: str | None = None,
) -> SuiteAuditReport:
    manifest_path = Path(manifest).resolve()
    try:
        raw = yaml.safe_load(manifest_path.read_text(encoding="utf-8")) or {}
    except (OSError, yaml.YAMLError) as exc:
        raise SuiteError(f"无法读取套件清单 {manifest_path}: {exc}") from exc
    root = _mapping(raw, "套件清单")
    suite = _mapping(root.get("suite", {}), "suite")
    selected_scale = str(scale or suite.get("scale") or "").strip().lower().replace("_", "-")
    if selected_scale not in VALID_SCALES:
        raise SuiteError(f"suite.scale 必须是 {', '.join(VALID_SCALES)}")
    report = SuiteAuditReport(manifest_path, selected_scale)
    policy = load_volume_policy()
    expected_values = suite.get("required_documents")
    if expected_values is None:
        expected_values = [item.code for item in iter_document_types()]
    if not isinstance(expected_values, list) or not expected_values:
        raise SuiteError("suite.required_documents 必须是非空列表")
    expected: list[str] = []
    for value in expected_values:
        code = get_document_type(str(value)).code
        if code in expected:
            _add(report, "ERROR", "SUITE_REQUIRED_DUPLICATE", f"required_documents 重复：{code}", code)
        else:
            expected.append(code)

    entries = root.get("documents")
    if not isinstance(entries, list) or not entries:
        raise SuiteError("documents 必须是非空列表")
    by_code: dict[str, dict[str, Any]] = {}
    for index, value in enumerate(entries, 1):
        entry = _mapping(value, f"documents[{index}]")
        code = get_document_type(str(entry.get("type", ""))).code
        if code in by_code:
            _add(report, "ERROR", "SUITE_DOCUMENT_DUPLICATE", f"documents 中 {code} 重复", code)
        else:
            by_code[code] = entry

    for code in expected:
        if code not in by_code:
            _add(report, "ERROR", "SUITE_DOCUMENT_MISSING", f"缺少必需文档 {code}", code)
    for code in sorted(set(by_code) - set(expected)):
        _add(report, "WARN", "SUITE_DOCUMENT_EXTRA", f"清单包含未声明为必需的文档 {code}", code)

    base = manifest_path.parent
    for code in expected:
        entry = by_code.get(code)
        if entry is None:
            continue
        try:
            markdown_path = _path(base, entry.get("markdown"), f"{code}.markdown")
            docx_path = _path(base, entry.get("docx"), f"{code}.docx")
            volume_path = _path(base, entry.get("volume"), f"{code}.volume")
        except SuiteError as exc:
            _add(report, "ERROR", "SUITE_PATH_INVALID", str(exc), code)
            continue
        item = SuiteDocumentResult(
            code, str(markdown_path), str(docx_path), str(volume_path)
        )
        report.documents.append(item)
        missing = [path for path in (markdown_path, docx_path, volume_path) if not path.is_file()]
        if missing:
            for path in missing:
                _add(report, "ERROR", "SUITE_FILE_MISSING", f"找不到 {path}", code)
            continue

        baseline: Path | None = None
        if entry.get("baseline"):
            try:
                baseline = _path(base, entry.get("baseline"), f"{code}.baseline")
            except SuiteError as exc:
                _add(report, "ERROR", "SUITE_BASELINE_INVALID", str(exc), code)
        elif code == "SDD" and "SRS" in by_code:
            baseline = _path(base, by_code["SRS"].get("markdown"), "SRS.markdown")
        elif code == "SSDD" and "SSS" in by_code:
            baseline = _path(base, by_code["SSS"].get("markdown"), "SSS.markdown")

        content_report = audit_markdown_with_profile(
            markdown_path,
            profile_name="release",
            document_type=code,
            baseline_srs=baseline,
            scale=selected_scale,
        )
        if not content_report.passed:
            _add(
                report,
                "ERROR",
                "SUITE_CONTENT_AUDIT_FAILED",
                content_report.to_text(),
                code,
            )

        docx_report = audit_docx(docx_path, profile="release")
        if not getattr(docx_report, "passed", False):
            _add(
                report,
                "ERROR",
                "SUITE_DOCX_AUDIT_FAILED",
                docx_report.to_text(),
                code,
            )

        try:
            volume = _volume_payload(volume_path)
        except SuiteError as exc:
            _add(report, "ERROR", "SUITE_VOLUME_INVALID", str(exc), code)
            continue
        reported_code = str(volume.get("document_type", "")).upper()
        reported_scale = str(volume.get("scale", "")).lower()
        if reported_code != code:
            _add(report, "ERROR", "SUITE_VOLUME_TYPE_MISMATCH", f"体量报告声明 {reported_code}", code)
        if reported_scale != selected_scale:
            _add(report, "ERROR", "SUITE_VOLUME_SCALE_MISMATCH", f"体量报告规模为 {reported_scale}", code)
        if not volume.get("passed"):
            _add(report, "ERROR", "SUITE_VOLUME_NOT_PASSED", "单文档体量报告未通过", code)

        document = parse_markdown(markdown_path)
        source_digest = sha256_text(document.raw)
        docx_digest = sha256_file(docx_path)
        if str(volume.get("source_sha256", "")) != source_digest:
            _add(report, "ERROR", "SUITE_SOURCE_HASH_MISMATCH", "Markdown 与体量报告未绑定或已被修改", code)
        if str(volume.get("docx_sha256", "")) != docx_digest:
            _add(report, "ERROR", "SUITE_DOCX_HASH_MISMATCH", "DOCX 与体量报告未绑定或已被修改", code)
        if int(volume.get("policy_version", -1)) != int(policy.get("policy_version", 1)):
            _add(report, "ERROR", "SUITE_POLICY_VERSION_MISMATCH", "体量报告使用了不同策略版本", code)

        pages = int(volume.get("pages", 0))
        floor = minimum_pages(code, selected_scale, entry.get("min_pages"))
        units = int(volume.get("effective_units", 0))
        unit_floor = int(volume.get("minimum_effective_units", 0))
        item.pages = pages
        item.minimum_pages = floor
        if pages < floor:
            _add(report, "ERROR", "SUITE_PAGE_FLOOR", f"页数 {pages} 低于下限 {floor}", code)
        if units < unit_floor:
            _add(report, "ERROR", "SUITE_EFFECTIVE_UNITS", f"等效内容单位 {units} 低于 {unit_floor}", code)
        item.passed = not any(
            issue.severity == "ERROR" and issue.document_type == code
            for issue in report.issues
        )
        report.total_pages += pages

    override = suite.get("min_portfolio_pages")
    report.minimum_total_pages = (
        int(override)
        if override is not None
        else int(policy["portfolio_min_pages"][selected_scale])
    )
    if report.total_pages < report.minimum_total_pages:
        _add(
            report,
            "ERROR",
            "SUITE_PORTFOLIO_PAGE_FLOOR",
            f"套件总页数 {report.total_pages} 低于 {selected_scale} 下限 {report.minimum_total_pages}",
        )
    return report
'''
(PKG / "suite.py").write_text(suite_py, encoding="utf-8")

cli = PKG / "cli.py"
text = cli.read_text(encoding="utf-8")
if "from .suite import" not in text:
    anchor = "from .volume import VolumeError, audit_rendered_volume, load_volume_policy, resolve_scale\n"
    assert anchor in text
    text = text.replace(anchor, anchor + "from .suite import SuiteError, audit_suite_manifest\n")
if 'sub.add_parser("audit-suite"' not in text:
    marker = '    policy = sub.add_parser("volume-policy", help="查看某文档类型/规模的工程体量门禁（非标准条款）")\n'
    assert marker in text
    block = '''    suite = sub.add_parser("audit-suite", help="审核一组已生成的 Markdown、DOCX 和体量报告")
    suite.add_argument("manifest")
    suite.add_argument("--scale", choices=["prototype", "small", "medium", "large", "very-large"])
    suite.add_argument("--json")

'''
    text = text.replace(marker, block + marker)
if 'args.command == "audit-suite"' not in text:
    marker = '        if args.command == "volume-policy":\n'
    assert marker in text
    block = '''        if args.command == "audit-suite":
            report = audit_suite_manifest(args.manifest, scale=args.scale)
            if args.json:
                Path(args.json).write_text(report.to_json(), encoding="utf-8")
            print(report.to_text())
            return 0 if report.passed else 4

'''
    text = text.replace(marker, block + marker)
text = text.replace(
    "except (ValueError, FileNotFoundError, FrontMatterError, RenderError, ImportWordError, FinalizeError, VolumeError) as exc:",
    "except (ValueError, FileNotFoundError, FrontMatterError, RenderError, ImportWordError, FinalizeError, VolumeError, SuiteError) as exc:",
)
cli.write_text(text, encoding="utf-8")

(SKILL / "examples/suite.example.yaml").write_text('''suite:
  name: 示例软件文档套件
  scale: prototype
  required_documents: [SRS, SDD]
  # 完整项目默认应列出经裁剪确认的全部适用文档。
  min_portfolio_pages: 2

documents:
  - type: SRS
    markdown: SRS.example.md
    docx: ../../dist/SRS.docx
    volume: ../../dist/SRS.volume.json
  - type: SDD
    markdown: SDD.example.md
    docx: ../../dist/SDD.docx
    volume: ../../dist/SDD.volume.json
    baseline: SRS.example.md
''', encoding="utf-8")

(ROOT / "docs/PAGE-COUNT-RESEARCH.md").write_text('''# GJB 438C 文档页数调研与采用口径

## 调研结论

公开网络中可以找到标准介绍、模板、工具和零散截图片段，但没有形成可信、可复核的“二十类完整生产项目文档及页数”样本集。公开 GitHub 代码检索也没有找到可直接作为统计样本的完整 GJB 438C DOCX 成品。因此，本仓库不声称“标准规定所有文档至少几百页”，也不根据来源不明的网盘材料制定硬阈值。

仓库内二十份 DOCX 是章节模板或示例壳。它们的真实渲染页数会由 `tools/report_template_pages.py` 写入 `TEMPLATE-PAGE-BASELINE.md`。这些页数只能证明模板骨架规模，不能代表填入真实需求、设计、接口、数据库和测试证据后的交付规模。

## 工程判断

大型或超大型系统中，SRS、SSS、SDD、SSDD、STD、STR 等核心规格、设计和测试文档达到数百页是合理且常见的结果；测试说明甚至可能因为大量测试用例达到更高体量。SIP、STrP、SVD 等职责较窄的文件通常不应被强制扩成数百页。

因此采用五档项目规模和“单文档 + 全套”双层门禁。页数只是一项外观结果，还必须与章节覆盖、可见正文、来源证据、需求、设计单元、接口、数据、测试用例及追踪项的数量共同通过。来源材料不够时阻断发布，不允许注水。
''', encoding="utf-8")

report_tool = r'''#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path
import re
import shutil
import subprocess
import tempfile

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_ROOT = ROOT / "GJB438C全套模版" / "438C-2021全套模板"
CODE_RE = re.compile(r"\[(\d+)\]\[([^]]+)\]\s*(.+?)-438C-2021\.docx$")


def office() -> str:
    value = shutil.which("libreoffice") or shutil.which("soffice")
    if not value:
        raise SystemExit("LibreOffice/soffice is required")
    return value


def pages(path: Path) -> int:
    with tempfile.TemporaryDirectory(prefix="gjb438c-template-pages-") as name:
        out = Path(name)
        result = subprocess.run(
            [office(), "--headless", "--convert-to", "pdf", "--outdir", str(out), str(path.resolve())],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=180,
        )
        pdf = out / f"{path.stem}.pdf"
        if result.returncode or not pdf.is_file():
            raise RuntimeError(result.stderr.strip() or result.stdout.strip())
        return len(PdfReader(str(pdf)).pages)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", default="docs/TEMPLATE-PAGE-BASELINE.md")
    parser.add_argument("--json", default="docs/template-page-baseline.json")
    args = parser.parse_args()
    rows = []
    for path in sorted(TEMPLATE_ROOT.glob("*.docx")):
        match = CODE_RE.search(path.name)
        if not match:
            continue
        document = Document(path)
        rows.append(
            {
                "number": int(match.group(1)),
                "code": match.group(2),
                "name": match.group(3),
                "filename": path.name,
                "rendered_pages": pages(path),
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
            }
        )
    if len(rows) != 20:
        raise SystemExit(f"expected 20 templates, got {len(rows)}")
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# 二十份源模板真实渲染页数基线",
        "",
        "> 这是模板/示例壳的页数，不是生产项目成品下限，也不是 GJB 438C 的统一页数条款。",
        "",
        "| 编号 | 类型 | 文档 | 模板页数 | 段落 | 表格 |",
        "|---:|---|---|---:|---:|---:|",
    ]
    for row in rows:
        lines.append(
            f"| {row['number']} | {row['code']} | {row['name']} | "
            f"{row['rendered_pages']} | {row['paragraphs']} | {row['tables']} |"
        )
    total = sum(row["rendered_pages"] for row in rows)
    lines += [
        "",
        f"模板合计页数：**{total}**。",
        "",
        "这些模板用于校准章节和初始版式。真实项目页数必须由项目规模、适用性裁剪和可核查技术证据共同决定。",
    ]
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")
    json_path = Path(args.json)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
'''
(ROOT / "tools/report_template_pages.py").write_text(report_tool, encoding="utf-8")

volume_docs = ROOT / "docs/VOLUME-POLICY.md"
current = volume_docs.read_text(encoding="utf-8")
if "## 全套文档审计" not in current:
    current += '''

## 全套文档审计

每次 `render --volume-json` 都会把 Markdown 和 DOCX 的 SHA-256、策略版本、真实渲染页数及等效内容单位写入体量报告。全套发布使用 manifest 将三者绑定：

```bash
gjb438c audit-suite suite.yaml --scale large --json dist/suite.audit.json
```

`audit-suite` 会重新执行 Markdown 发布审计和 DOCX 格式审计，校验体量报告哈希，检查每类文档页数，并对经裁剪确认的全套文档执行总页数门禁。修改 Word 或 Markdown 后，旧体量报告会因为哈希不一致而失效，必须重新生成。
'''
    volume_docs.write_text(current, encoding="utf-8")

root_readme = ROOT / "README.md"
current = root_readme.read_text(encoding="utf-8")
if "gjb438c audit-suite" not in current:
    current = current.replace(
        "gjb438c volume-policy --type SDD --scale large\n```",
        "gjb438c volume-policy --type SDD --scale large\n\n# 汇总审核一套已经生成的文档\ngjb438c audit-suite suite.yaml --scale large --json dist/suite.audit.json\n```",
    )
    current += "\n模板真实渲染页数基线见 `docs/TEMPLATE-PAGE-BASELINE.md`；调研边界见 `docs/PAGE-COUNT-RESEARCH.md`。\n"
    root_readme.write_text(current, encoding="utf-8")

skill_readme = SKILL / "README.md"
current = skill_readme.read_text(encoding="utf-8")
if "audit-suite" not in current:
    current += '''

## 全套发布

单文档渲染时保存 `--volume-json`，然后通过 suite manifest 绑定 Markdown、DOCX 和体量报告：

```bash
gjb438c audit-suite examples/suite.example.yaml --scale prototype
```

生产清单应显式列出经过适用性裁剪确认的文档类型。完整二十类项目默认按体量策略中的 `portfolio_min_pages` 检查总页数。
'''
    skill_readme.write_text(current, encoding="utf-8")

(TESTS / "test_volume_hash_binding.py").write_text('''from pathlib import Path
from gjb438c_suite.markdown_doc import parse_markdown
from gjb438c_suite.profiles import load_profile
from gjb438c_suite import volume


def test_volume_report_binds_markdown_and_docx(monkeypatch, tmp_path: Path):
    source = Path(__file__).resolve().parents[1] / "examples/SRS.example.md"
    document = parse_markdown(source)
    docx = tmp_path / "SRS.docx"
    docx.write_bytes(b"docx-placeholder")
    monkeypatch.setattr(volume, "rendered_page_count", lambda _path: 3)
    result = volume.audit_rendered_volume(
        document, load_profile("SRS"), docx, scale="prototype", min_pages_override=1
    )
    assert result.passed
    assert result.source_sha256 == volume.sha256_text(document.raw)
    assert result.docx_sha256 == volume.sha256_file(docx)
    assert result.policy_version >= 1
''', encoding="utf-8")

(TESTS / "test_suite_audit.py").write_text('''from pathlib import Path
import json
import shutil
import yaml

from gjb438c_suite import suite
from gjb438c_suite.markdown_doc import parse_markdown
from gjb438c_suite.volume import load_volume_policy, sha256_file, sha256_text


class Passed:
    passed = True
    def to_text(self):
        return "PASS"


def test_suite_audit_verifies_hashes_and_portfolio(monkeypatch, tmp_path: Path):
    source = Path(__file__).resolve().parents[1] / "examples/SRS.example.md"
    markdown = tmp_path / "SRS.md"
    shutil.copy2(source, markdown)
    docx = tmp_path / "SRS.docx"
    docx.write_bytes(b"synthetic-docx-for-suite-test")
    payload = {
        "document_type": "SRS",
        "scale": "prototype",
        "pages": 3,
        "minimum_pages": 1,
        "visible_chars": 1000,
        "effective_units": 1000,
        "minimum_effective_units": 260,
        "passed": True,
        "issues": [],
        "policy_version": load_volume_policy()["policy_version"],
        "source_sha256": sha256_text(parse_markdown(markdown).raw),
        "docx_sha256": sha256_file(docx),
    }
    volume = tmp_path / "SRS.volume.json"
    volume.write_text(json.dumps(payload), encoding="utf-8")
    manifest = tmp_path / "suite.yaml"
    manifest.write_text(yaml.safe_dump({
        "suite": {"scale": "prototype", "required_documents": ["SRS"], "min_portfolio_pages": 1},
        "documents": [{"type": "SRS", "markdown": "SRS.md", "docx": "SRS.docx", "volume": "SRS.volume.json"}],
    }, allow_unicode=True), encoding="utf-8")
    monkeypatch.setattr(suite, "audit_markdown_with_profile", lambda *a, **k: Passed())
    monkeypatch.setattr(suite, "audit_docx", lambda *a, **k: Passed())
    report = suite.audit_suite_manifest(manifest)
    assert report.passed
    assert report.total_pages == 3
    docx.write_bytes(b"changed-after-report")
    report = suite.audit_suite_manifest(manifest)
    assert not report.passed
    assert any(issue.code == "SUITE_DOCX_HASH_MISMATCH" for issue in report.errors)
''', encoding="utf-8")

workflow = ROOT / ".github/workflows/gjb438c-md-first.yml"
current = workflow.read_text(encoding="utf-8")
if "Suite audit CLI availability" not in current:
    current = current.replace(
        "      - name: Content gates\n",
        "      - name: Suite audit CLI availability\n        run: gjb438c audit-suite --help\n      - name: Content gates\n",
    )
    workflow.write_text(current, encoding="utf-8")

print("suite audit integration staged")
