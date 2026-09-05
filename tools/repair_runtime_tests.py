from __future__ import annotations

from pathlib import Path
import re
import textwrap

ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "skills" / "gjb438c-md-first" / "gjb438c_suite"
TESTS = ROOT / "skills" / "gjb438c-md-first" / "tests"
TESTS.mkdir(parents=True, exist_ok=True)


def write(relative: str, content: str) -> None:
    path = ROOT / relative
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(textwrap.dedent(content).lstrip(), encoding="utf-8")


# Harden tier normalization and make invalid per-document suite overrides a
# reportable failure instead of aborting the entire suite audit.
profile_quality = PKG / "profile_quality.py"
text = profile_quality.read_text(encoding="utf-8")
text = text.replace(
    'key = str(value or "prototype").strip().lower().replace(" ", "-")',
    'key = str(value or "prototype").strip().lower().replace(" ", "-").replace("_", "-")',
)
profile_quality.write_text(text, encoding="utf-8")

suite_path = PKG / "suite.py"
text = suite_path.read_text(encoding="utf-8")
old = '''            floor = minimum_body_pages(code, normalized_tier, entry.get("min_body_pages"))
            total_floor += floor
'''
new = '''            try:
                floor = minimum_body_pages(code, normalized_tier, entry.get("min_body_pages"))
            except VolumeError as exc:
                _add(report, "ERROR", "SUITE_PAGE_OVERRIDE_INVALID", str(exc), code)
                floor = minimum_body_pages(code, normalized_tier)
            total_floor += floor
'''
if old not in text:
    raise SystemExit("suite floor block not found")
suite_path.write_text(text.replace(old, new), encoding="utf-8")

write(
    "skills/gjb438c-md-first/tests/test_runtime_wiring.py",
    r'''
    from pathlib import Path

    import pytest
    import yaml

    from gjb438c_suite.cli import build_parser, main
    from gjb438c_suite.profile_quality import (
        audit_markdown_with_profile,
        load_profile_mapping,
        tier_minimum,
    )
    from gjb438c_suite.profiles import load_profile
    from gjb438c_suite.registry import iter_document_types
    from gjb438c_suite.suite import initialize_suite
    from gjb438c_suite.volume import VolumeError, minimum_body_pages


    CODES = [item.code for item in iter_document_types()]


    def test_cli_exposes_real_profile_volume_and_suite_commands():
        parser = build_parser()
        for argv in (
            ["profile", "--type", "SRS"],
            ["volume-policy", "--type", "SDD", "--scale", "large"],
            ["suite-init", "--output", "out"],
            ["audit-volume", "x.docx", "--source", "x.md"],
            ["audit-suite", "suite.yaml"],
        ):
            parser.parse_args(argv)


    def test_all_20_profiles_are_loadable_and_isolated():
        assert len(CODES) == 20
        for code in CODES:
            first = load_profile(code)
            second = load_profile(code)
            assert first is not second
            mapping = load_profile_mapping(code)
            assert mapping["code"] == code
            assert mapping["outline"]
            assert mapping["artifact_contracts"]


    def test_tier_specific_minimum_is_not_collapsed_to_one():
        mapping = load_profile_mapping("SRS")
        requirement = next(item for item in mapping["artifact_contracts"] if item["kind"] == "requirement")
        assert tier_minimum(requirement["minimum"], "large") == int(requirement["minimum"]["large"])
        assert tier_minimum(requirement["minimum"], "large") > 1


    def test_init_works_without_repository_template_root(tmp_path: Path):
        project = Path(__file__).resolve().parents[1] / "examples/project.yaml"
        for code in CODES:
            target = tmp_path / f"{code}.md"
            assert main(["init", "--type", code, "--project", str(project), "--output", str(target)]) == 0
            raw = target.read_text(encoding="utf-8")
            assert raw.startswith("---")
            assert code in raw


    def test_suite_init_creates_native_unicode_20_document_workspace(tmp_path: Path):
        project = Path(__file__).resolve().parents[1] / "examples/project.yaml"
        manifest = initialize_suite(
            tmp_path / "suite",
            project=project,
            tier="large",
            min_body_pages=300,
            suite_id="PUBLIC-DEMO",
        )
        payload = yaml.safe_load(manifest.read_text(encoding="utf-8"))
        assert len(payload["documents"]) == 20
        assert set(payload["documents"]) == set(CODES)
        for code, entry in payload["documents"].items():
            path = manifest.parent / entry["markdown"]
            assert path.is_file()
            assert "#U" not in path.name
            metadata = yaml.safe_load(path.read_text(encoding="utf-8").split("---", 2)[1])
            assert metadata["document"]["type"] == code
            assert metadata["document"]["status"] == "draft"
            assert metadata["quality"]["min_body_pages"] >= 300


    def test_page_override_can_raise_but_not_lower_profile_floor():
        floor = minimum_body_pages("SSS", "large")
        with pytest.raises(VolumeError):
            minimum_body_pages("SSS", "large", floor - 1)
        assert minimum_body_pages("SSS", "large", floor + 50) == floor + 50


    def test_zcode_style_plain_markdown_is_rejected(tmp_path: Path):
        source = tmp_path / "SSS.md"
        source.write_text("# 修订说明\n\n# 1 范围\n\n这只是普通摘要，没有元数据和结构化证据。\n", encoding="utf-8")
        report = audit_markdown_with_profile(
            source,
            profile="release",
            document_type="SSS",
            tier="large",
        )
        assert not report.passed
        text = report.to_text()
        assert "ERROR" in text
        assert "gjb-requirement" in text or "元数据" in text or "front" in text.lower()
    ''',
)

write(
    "skills/gjb438c-md-first/tests/test_volume_hardening.py",
    r'''
    from pathlib import Path

    from gjb438c_suite.markdown_doc import parse_markdown
    from gjb438c_suite.volume import duplicate_prose_metrics, markdown_volume_issues


    def _source(tmp_path: Path, paragraphs: list[str]) -> Path:
        path = tmp_path / "SRS.md"
        path.write_text(
            """---
    document:
      type: SRS
      id: TEST-SRS
      status: approved
    quality:
      tier: large
    software:
      name: 示例软件
      identifier: TEST
      version: V1.0
    sources:
      - id: SRC-1
        title: 示例来源
    ---
    # 1 范围

    """ + "\n\n".join(paragraphs),
            encoding="utf-8",
        )
        return path


    def test_repeated_long_prose_is_not_counted_as_real_volume(tmp_path: Path):
        repeated = (
            "本段用于验证重复内容门禁，真实文档必须给出项目特定的输入、处理、输出、异常、"
            "边界条件、来源和验证方法，机械复制同一段落不能增加有效正文体量。"
        )
        document = parse_markdown(_source(tmp_path, [repeated] * 20))
        duplicate_chars, ratio, repeat = duplicate_prose_metrics(document)
        assert duplicate_chars > 0
        assert ratio > 0.8
        assert repeat == 20
        issues = markdown_volume_issues(document, "SRS", "large", "release")
        assert any(item["code"] == "VOLUME_DUPLICATE_CONTENT" for item in issues)


    def test_distinct_prose_is_not_marked_as_duplicate(tmp_path: Path):
        paragraphs = [
            f"第{index}段给出不同的条件、数据、状态、处理路径、异常恢复和验收结果。"
            + (f"唯一证据编号为 EVID-{index:03d}。" * 8)
            for index in range(20)
        ]
        document = parse_markdown(_source(tmp_path, paragraphs))
        duplicate_chars, ratio, repeat = duplicate_prose_metrics(document)
        assert duplicate_chars == 0
        assert ratio == 0
        assert repeat <= 1
    ''',
)

write(
    "skills/gjb438c-md-first/tests/test_evidence_appendix.py",
    r'''
    from pathlib import Path

    from docx import Document

    from gjb438c_suite.evidence import append_evidence_appendix
    from gjb438c_suite.markdown_doc import parse_markdown
    from gjb438c_suite.profile_quality import artifact_mapping


    def test_structured_evidence_is_visible_in_word(tmp_path: Path):
        source = Path(__file__).resolve().parents[1] / "examples/SRS.example.md"
        parsed = parse_markdown(source)
        artifacts = list(parsed.artifacts)
        assert artifacts
        first_id = str(artifact_mapping(artifacts[0]).get("id", ""))
        assert first_id
        docx = tmp_path / "evidence.docx"
        Document().save(docx)
        append_evidence_appendix(docx, parsed)
        word = Document(docx)
        visible = "\n".join(
            [paragraph.text for paragraph in word.paragraphs]
            + [cell.text for table in word.tables for row in table.rows for cell in row.cells]
        )
        assert "结构化工程证据" in visible
        assert first_id in visible
    ''',
)

write(
    "skills/gjb438c-md-first/tests/test_suite_guards.py",
    r'''
    from pathlib import Path

    import pytest
    import yaml

    from gjb438c_suite.suite import SuiteError, audit_suite_manifest, initialize_suite


    def test_suite_init_refuses_nonempty_directory(tmp_path: Path):
        target = tmp_path / "suite"
        target.mkdir()
        (target / "keep.txt").write_text("do not overwrite", encoding="utf-8")
        with pytest.raises(SuiteError):
            initialize_suite(target)
        assert (target / "keep.txt").read_text(encoding="utf-8") == "do not overwrite"


    def test_suite_audit_reports_lowered_document_floor(tmp_path: Path):
        manifest = initialize_suite(tmp_path / "suite", tier="large")
        data = yaml.safe_load(manifest.read_text(encoding="utf-8"))
        data["documents"]["SSS"]["min_body_pages"] = 1
        manifest.write_text(yaml.safe_dump(data, allow_unicode=True, sort_keys=False), encoding="utf-8")
        report = audit_suite_manifest(manifest, audit_profile="review")
        assert not report.passed
        assert any(item.code == "SUITE_PAGE_OVERRIDE_INVALID" for item in report.issues)
    ''',
)

write(
    "tools/stress_volume_gate.py",
    r'''
    from __future__ import annotations

    from pathlib import Path
    import tempfile

    from docx import Document

    from gjb438c_suite.volume import audit_rendered_volume


    def markdown(path: Path, paragraphs: list[str]) -> None:
        path.write_text(
            """---
    document:
      type: SRS
      id: STRESS-SRS
      status: approved
    quality:
      tier: large
    software:
      name: 公开压力测试软件
      identifier: STRESS
      version: V1.0
    sources:
      - id: SRC-STRESS
        title: 压力测试数据
    ---
    # 1 范围

    """ + "\n\n".join(paragraphs),
            encoding="utf-8",
        )


    def word(path: Path, paragraphs: list[str]) -> None:
        doc = Document()
        for index, title in enumerate(("首页", "签字页", "修改页", "目录")):
            doc.add_paragraph(title)
            doc.add_page_break()
        for index, value in enumerate(paragraphs):
            doc.add_heading(f"1.{index + 1} 压力测试条目 {index + 1}", level=2)
            doc.add_paragraph(value)
            if index + 1 < len(paragraphs):
                doc.add_page_break()
        doc.save(path)


    def main() -> int:
        with tempfile.TemporaryDirectory(prefix="gjb438c-stress-") as name:
            root = Path(name)
            unique = [
                (f"第{index + 1}页描述唯一的项目条件、输入、处理、输出、状态转换、异常恢复、"
                 f"验证方法和来源证据 EVID-{index + 1:04d}。" +
                 (f"本页唯一序列为 {index + 1:04d}，不得与其他页面替换。" * 10))
                for index in range(300)
            ]
            repeated = [
                "本页是机械重复内容，虽然通过分页符形成很多页，但没有新增需求、设计、接口、"
                "测试或来源证据，因此必须被重复页和重复长段落门禁拒绝。" * 12
            ] * 300
            unique_md = root / "unique.md"
            unique_docx = root / "unique.docx"
            repeated_md = root / "repeated.md"
            repeated_docx = root / "repeated.docx"
            markdown(unique_md, unique)
            markdown(repeated_md, repeated)
            word(unique_docx, unique)
            word(repeated_docx, repeated)
            good = audit_rendered_volume(
                unique_md,
                "SRS",
                unique_docx,
                tier="large",
                min_body_pages_override=300,
                body_start_page=5,
            )
            bad = audit_rendered_volume(
                repeated_md,
                "SRS",
                repeated_docx,
                tier="large",
                min_body_pages_override=300,
                body_start_page=5,
            )
            print(good.to_text())
            print(bad.to_text())
            if not good.passed:
                raise SystemExit("unique 300-page stress document was rejected")
            if bad.passed:
                raise SystemExit("repeated 300-page padding document was accepted")
        return 0


    if __name__ == "__main__":
        raise SystemExit(main())
    ''',
)

write(
    ".github/workflows/gjb438c-md-first.yml",
    r'''
    name: gjb438c-md-first

    on:
      pull_request:
        paths:
          - 'skills/gjb438c-*/**'
          - 'legacy-skills/**'
          - 'tools/stress_volume_gate.py'
          - '.github/workflows/gjb438c-md-first.yml'
          - 'README.md'
      push:
        branches: [main, fix/wire-20-profiles-volume-suite-20260905]
        paths:
          - 'skills/gjb438c-*/**'
          - 'legacy-skills/**'
          - 'tools/stress_volume_gate.py'
          - '.github/workflows/gjb438c-md-first.yml'
          - 'README.md'
      workflow_dispatch:

    permissions:
      contents: read

    jobs:
      test:
        runs-on: ubuntu-latest
        env:
          PYTHONPATH: /usr/lib/python3/dist-packages
        steps:
          - uses: actions/checkout@v4
          - uses: actions/setup-python@v5
            with:
              python-version: '3.12'
              cache: pip
              cache-dependency-path: skills/gjb438c-md-first/pyproject.toml
          - name: Install Office rendering dependencies
            run: |
              sudo apt-get update
              sudo apt-get install -y --no-install-recommends libreoffice-writer python3-uno poppler-utils
          - name: Install suite
            run: python -m pip install -e 'skills/gjb438c-md-first[test]'
          - name: Compile and test
            run: |
              python -m compileall -q skills/gjb438c-md-first/gjb438c_suite
              python -m pytest skills/gjb438c-md-first/tests
          - name: Verify all 20 profiles are standalone
            run: |
              mkdir -p /tmp/gjb438c-all
              for type in SDP SIP STrP STP OCD SSS IRS SSDD IDD SRS SDD DBDD STD STR SPS SVD SUM CPM FSM SDSR; do
                gjb438c init --type "$type" \
                  --project skills/gjb438c-md-first/examples/project.yaml \
                  --output "/tmp/gjb438c-all/$type.md"
              done
              test "$(find /tmp/gjb438c-all -name '*.md' | wc -l)" -eq 20
              gjb438c profile --type STD --json >/tmp/STD.profile.json
              gjb438c volume-policy --type SDD --scale large --json >/tmp/SDD.volume.json
          - name: Verify suite workspace generation
            run: |
              gjb438c suite-init \
                --project skills/gjb438c-md-first/examples/project.yaml \
                --output /tmp/gjb438c-suite \
                --scale large \
                --min-body-pages 300
              test -f /tmp/gjb438c-suite/suite.yaml
              test "$(find /tmp/gjb438c-suite/docs -name '*.md' | wc -l)" -eq 20
              ! find /tmp/gjb438c-suite -name '*#U*' -print -quit | grep .
          - name: Ensure integration scaffolding is absent
            run: |
              test ! -e .bootstrap-v03
              test ! -e tools/apply_v03_profiles.py
              test ! -e .github/workflows/apply-v03-profiles.yml
              test ! -e .github/workflows/finalize-v03.yml

      volume-stress:
        needs: test
        runs-on: ubuntu-latest
        timeout-minutes: 15
        env:
          PYTHONPATH: /usr/lib/python3/dist-packages
        steps:
          - uses: actions/checkout@v4
          - uses: actions/setup-python@v5
            with:
              python-version: '3.12'
          - name: Install Office rendering dependencies
            run: |
              sudo apt-get update
              sudo apt-get install -y --no-install-recommends libreoffice-writer python3-uno poppler-utils
          - name: Install suite
            run: python -m pip install -e 'skills/gjb438c-md-first[test]'
          - name: Prove real 300-page content passes and padding fails
            run: python tools/stress_volume_gate.py
    ''',
)

readme = ROOT / "README.md"
text = readme.read_text(encoding="utf-8")
section = r'''

## 20 类 Profile、体量门禁与整套审计

默认安装包含一个共享核心和 20 个文档专用路由 Skill。`gjb438c init` 默认读取安装包内置 Profile，不依赖仓库根目录的 DOCX 模板。

```bash
gjb438c profile --type SSS
gjb438c suite-init --project project.yaml --output project-docs \
  --scale large --min-body-pages 300
```

每份文档必须先通过 Markdown 内容/Profile 审计，再原子生成 Word：

```bash
gjb438c audit docs/SSS.md --profile release --scale large
gjb438c render docs/SSS.md --output dist/SSS.docx \
  --profile release --scale large --min-body-pages 300 \
  --refresh-toc --volume-json reports/SSS.volume.json
gjb438c audit-suite suite.yaml --profile release --write-volume-reports
```

体量下限是本仓库和具体项目的工程质量策略，不冒充 GJB 438C 的统一页数条款。正文页数不含首页、签字页、修改页和目录；重复段落、重复页、薄页、隐藏证据块和空白页不能计作有效体量。调用方只能提高下限，不能降低 Profile 策略。
'''
if "## 20 类 Profile、体量门禁与整套审计" not in text:
    readme.write_text(text.rstrip() + section, encoding="utf-8")

write(
    "docs/RUNTIME-AND-VOLUME-GATES.md",
    r'''
    # GJB 438C 文档运行时与体量门禁

    本仓库使用一个共享 Markdown-first 核心和 20 个文档 Profile。Profile 包含章节结构、来源材料、结构化证据字段、基线关系和按项目规模划分的工程下限。

    ## 发布条件

    正式 Word 必须同时满足：

    1. Markdown front matter、来源清单和文档状态完整；
    2. 文档类型 Profile 的章节、字段、条目数量和稳定标识完整；
    3. 上下游基线存在，结构化引用可以解析；
    4. Word 字体、字号、行距、目录、页码和前三页通过审计；
    5. LibreOffice 实际渲染后的正文页数达到下限；
    6. 可见正文字符、薄页比例、重复页比例和重复长段落比例通过门禁；
    7. 体量报告中的 Markdown 与 DOCX SHA-256 与当前文件一致；
    8. 整套 `suite.yaml` 审计通过。

    `--min-body-pages` 只能提高下限。`prototype` 只能用于带 `quality.fixture: true` 的仓库测试夹具，不能作为生产发布档位。

    ## 为什么不只检查页数

    页数很容易被空白页、分页符、复制段落、放大图片和空表格伪造。因此发布门禁同时检查可见技术正文、结构化证据、跨文档追踪和实际页面内容。几百页只在真实需求、设计单元、接口、数据模型、操作步骤和测试用例足够多时才有意义。
    ''',
)

print("tests, CI and documentation written")
