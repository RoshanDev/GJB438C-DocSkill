"""Render structured evidence visibly, inside the source-tracked body bookmark."""
from __future__ import annotations
from collections import defaultdict
from pathlib import Path
import json

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .markdown_doc import MarkdownDocument, parse_markdown

TITLE = "附录 结构化工程证据"


def _text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list, tuple)):
        return json.dumps(value, ensure_ascii=False, indent=2, default=str)
    return str(value)


def append_evidence(document, source: MarkdownDocument, styles):
    if not source.artifacts:
        return None
    document.add_page_break()
    heading = document.add_paragraph(TITLE, style=styles['heading_1'])
    document.add_paragraph("以下记录来自 Markdown 的结构化证据。它们可供追踪核对，不能替代正文的技术论证或人工评审。", style=styles['body'])
    groups = defaultdict(list)
    for item in source.artifacts:
        groups[item.kind].append(item)
    for kind, items in sorted(groups.items()):
        document.add_paragraph(f"gjb-{kind}", style=styles['heading_2'])
        for item in items:
            document.add_paragraph(item.identifier or "未编号证据", style=styles['caption'])
            table = document.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for key, value in item.data.items():
                cells = table.add_row().cells
                cells[0].text, cells[1].text = str(key), _text(value)
                for cell in cells:
                    for p in cell.paragraphs:
                        p.style = styles['table']
                        p.paragraph_format.keep_together = False
                        p.paragraph_format.keep_with_next = False
            # One short table per artifact avoids quadratic cell traversal and
            # allows Word to split long rows rather than clipping them.
    return document.add_paragraph("结构化工程证据结束。", style=styles['body_no_indent'])


def append_evidence_appendix(docx: str | Path, source: str | Path | MarkdownDocument) -> Path:
    """Compatibility helper; the normal renderer already appends the evidence."""
    from .render import configure_styles, _patch_settings_with_source
    document = source if isinstance(source, MarkdownDocument) else parse_markdown(source)
    target = Path(docx)
    word = Document(target)
    if any(p.text == TITLE for p in word.paragraphs):
        return target
    ends = word.element.xpath('.//w:bookmarkEnd[@w:id="1000"]')
    last = append_evidence(word, document, configure_styles(word))
    if last is not None and ends:
        end = ends[0]
        end.getparent().remove(end)
        last._p.append(end)
    word.save(target)
    _patch_settings_with_source(target, document.raw)
    return target
