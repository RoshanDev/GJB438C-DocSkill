from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from gjb438c_suite.markdown_doc import extract_template_outline, render_skeleton
from gjb438c_suite.registry import get_document_type


def test_extract_outline_prefers_toc_and_renders_skeleton(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    for name in ("TOC 1", "TOC 2"):
        try:
            document.styles[name]
        except KeyError:
            document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("目录")
    document.add_paragraph("1 范围\t1", style="TOC 1")
    document.add_paragraph("1.1 标识\t1", style="TOC 2")
    document.add_paragraph("2 引用文档\t2", style="TOC 1")
    document.add_paragraph("正文开始")
    document.save(template)

    outline = extract_template_outline(template)
    assert [(item.level, item.number, item.title) for item in outline] == [
        (1, "1", "范围"),
        (2, "1.1", "标识"),
        (1, "2", "引用文档"),
    ]

    text = render_skeleton(
        document_type=get_document_type("SRS"),
        outline=outline,
        project={"software": {"name": "演示软件"}},
    )
    assert "document:\n  type: SRS" in text
    assert "# 1 范围" in text
    assert "## 1.1 标识" in text
    assert "```gjb-requirement" in text
