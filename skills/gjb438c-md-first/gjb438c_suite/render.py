from __future__ import annotations

import base64
from dataclasses import dataclass
from hashlib import sha256
import gzip
import json
from pathlib import Path
import re
import tempfile
from typing import Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

from .front_matter import FrontMatterError, render_front_matter
from .markdown_doc import MarkdownDocument, nested_get, parse_markdown
from .quality import AuditReport, audit_markdown
from .registry import default_front_matter_template, get_document_type

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}
BOOKMARK_NAME = "GJB_BODY"
DOCVAR_PREFIX = "GJB438C_SOURCE_"
DOCVAR_HASH = "GJB438C_BODY_TEXT_SHA256"
DOCVAR_SOURCE_HASH = "GJB438C_SOURCE_SHA256"
CAPTION_RE = re.compile(r"^(?:表|图)\s*[A-Za-z0-9一二三四五六七八九十附录.-]+(?:\s+|、).+")
HEADING_RE = re.compile(r"^(#{1,9})\s+(.+?)\s*$")
IMAGE_RE = re.compile(r"^!\[(?P<alt>[^]]*)\]\((?P<path>[^ )]+)(?:\s+\"(?P<title>[^\"]*)\")?\)\s*$")
LIST_RE = re.compile(r"^(?P<indent>\s*)(?P<mark>[-*+] |\d+[.)]\s+)(?P<text>.+)$")
TABLE_DIVIDER_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
LINK_RE = re.compile(r"\[([^]]+)\]\(([^)]+)\)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")


class RenderError(RuntimeError):
    pass


@dataclass(slots=True)
class RenderResult:
    output: Path
    markdown_report: AuditReport


def _set_style_fonts(style, east_asia: str, latin: str, size_pt: float, *, bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def _get_or_add_style(document: DocumentObject, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def _set_outline_level(style, level: int) -> None:
    ppr = style._element.get_or_add_pPr()
    for node in ppr.findall(qn("w:numPr")):
        ppr.remove(node)
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def configure_styles(document: DocumentObject) -> dict[str, str]:
    """Install the exact typography profile requested by the user."""
    # Do not modify the template's Normal style: the official first three
    # pages use it for spacing and signature lines. Global Normal changes can
    # push the signature block onto a fourth page. All generated content uses
    # dedicated GJB styles with explicit typography instead.
    normal = document.styles["Normal"]

    body = _get_or_add_style(document, "GJB 正文")
    body.base_style = normal
    _set_style_fonts(body, "宋体", "Times New Roman", 12)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    no_indent = _get_or_add_style(document, "GJB 正文无缩进")
    no_indent.base_style = body
    _set_style_fonts(no_indent, "宋体", "Times New Roman", 12)
    no_indent.paragraph_format.first_line_indent = Pt(0)

    caption = _get_or_add_style(document, "GJB 图表题")
    caption.base_style = normal
    _set_style_fonts(caption, "黑体", "Times New Roman", 10.5)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.keep_with_next = True

    table_text = _get_or_add_style(document, "GJB 表内文字")
    table_text.base_style = normal
    _set_style_fonts(table_text, "宋体", "Times New Roman", 10.5)
    table_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_text.paragraph_format.line_spacing = 1.0
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)

    code = _get_or_add_style(document, "GJB 代码")
    code.base_style = no_indent
    # The user's rule says all western text uses Times New Roman, including code.
    _set_style_fonts(code, "宋体", "Times New Roman", 10.5)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Pt(18)
    code.paragraph_format.right_indent = Pt(18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)

    toc_title = _get_or_add_style(document, "GJB 目录标题")
    toc_title.base_style = normal
    _set_style_fonts(toc_title, "宋体", "Times New Roman", 16)
    toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.line_spacing = 1.5
    toc_title.paragraph_format.space_before = Pt(0)
    toc_title.paragraph_format.space_after = Pt(18)

    heading_names: dict[str, str] = {}
    for level in range(1, 10):
        name = f"GJB 标题 {level}"
        style = _get_or_add_style(document, name)
        style.base_style = normal
        _set_style_fonts(style, "黑体", "Times New Roman", 12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        _set_outline_level(style, level)
        heading_names[str(level)] = name

    # Word-generated TOC entries use TOC 1...TOC 9. Modify only font,
    # line spacing and justification; preserve their level-specific indents.
    for level in range(1, 10):
        style = None
        for candidate in (f"TOC {level}", f"toc {level}"):
            try:
                style = document.styles[candidate]
                break
            except KeyError:
                pass
        if style is None:
            style = _get_or_add_style(document, f"TOC {level}")
        _set_style_fonts(style, "宋体", "Times New Roman", 12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    return {
        "body": body.name,
        "body_no_indent": no_indent.name,
        "caption": caption.name,
        "table": table_text.name,
        "code": code.name,
        "toc_title": toc_title.name,
        **{f"heading_{key}": value for key, value in heading_names.items()},
    }


def _set_section_a4(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)


def _set_page_numbering(section, *, fmt: str, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def _clear_footer(footer) -> None:
    for paragraph in footer.paragraphs:
        paragraph.clear()
    while len(footer.paragraphs) > 1:
        element = footer.paragraphs[-1]._element
        element.getparent().remove(element)


def _append_field(paragraph, instruction: str, cached_text: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate])
    if cached_text:
        paragraph.add_run(cached_text)
    paragraph.add_run()._r.append(end)


def _set_footer_page(section, *, roman: bool) -> None:
    section.footer.is_linked_to_previous = False
    _clear_footer(section.footer)
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if roman:
        _append_field(paragraph, " PAGE \\* ROMAN ")
    else:
        paragraph.add_run("— ")
        _append_field(paragraph, " PAGE ")
        paragraph.add_run(" —")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def _add_toc(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _append_field(
        paragraph,
        f' TOC \\o "1-9" \\h \\z \\u \\b "{BOOKMARK_NAME}" ',
        "请在 Word/WPS 中更新目录",
    )


def _bookmark_start(paragraph, bookmark_id: int = 1000) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), BOOKMARK_NAME)
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)


def _bookmark_end(paragraph, bookmark_id: int = 1000) -> None:
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _clean_inline(text: str) -> str:
    text = LINK_RE.sub(lambda match: f"{match.group(1)}（{match.group(2)}）", text)
    text = INLINE_CODE_RE.sub(r"\1", text)
    text = BOLD_RE.sub(r"\1", text)
    text = ITALIC_RE.sub(r"\1", text)
    return text.replace("<br>", "\n").replace("<br/>", "\n")


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [_clean_inline(cell.strip().replace("\\|", "|")) for cell in stripped.split("|")]


def _format_table(table, style_name: str) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style = style_name
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10.5)
                    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                    rfonts.set(qn("w:ascii"), "Times New Roman")
                    rfonts.set(qn("w:hAnsi"), "Times New Roman")
                    rfonts.set(qn("w:eastAsia"), "宋体")
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = tc_pr.find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    tc_pr.append(shading)
                shading.set(qn("w:fill"), "D9D2A9")


def _render_markdown_body(
    document: DocumentObject,
    markdown: MarkdownDocument,
    styles: dict[str, str],
    *,
    release: bool,
) -> tuple[object, object]:
    lines = markdown.visible_body.splitlines()
    doc_title = str(nested_get(markdown.metadata, "document.title", "")).strip()
    index = 0
    first_body_paragraph = None
    last_body_paragraph = None
    paragraph_buffer: list[str] = []
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    skipped_document_title = False
    chapter = "0"
    figure_counters: dict[str, int] = {}

    def remember(paragraph):
        nonlocal first_body_paragraph, last_body_paragraph
        if first_body_paragraph is None:
            first_body_paragraph = paragraph
        last_body_paragraph = paragraph
        return paragraph

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        value = _clean_inline(" ".join(part.strip() for part in paragraph_buffer if part.strip()))
        paragraph_buffer = []
        if not value:
            return
        style = styles["caption"] if CAPTION_RE.match(value) else styles["body"]
        paragraph = document.add_paragraph(value, style=style)
        remember(paragraph)

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if in_code:
            if stripped.startswith("```"):
                paragraph = document.add_paragraph(style=styles["code"])
                paragraph.add_run("\n".join(code_lines))
                remember(paragraph)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            in_code = True
            code_lang = stripped[3:].strip()
            index += 1
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 9)
            title = _clean_inline(heading_match.group(2).strip())
            if not skipped_document_title and title == doc_title:
                skipped_document_title = True
                index += 1
                continue
            skipped_document_title = True
            chapter_match = re.match(r"^(\d+)", title)
            if level == 1 and chapter_match:
                chapter = chapter_match.group(1)
            paragraph = document.add_paragraph(title, style=styles[f"heading_{level}"])
            remember(paragraph)
            index += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            flush_paragraph()
            image_path = Path(image_match.group("path"))
            if not image_path.is_absolute():
                image_path = (markdown.path.parent / image_path).resolve()
            if not image_path.is_file():
                if release:
                    raise RenderError(f"图片不存在：{image_path}")
                paragraph = document.add_paragraph(f"[图片缺失：{image_path}]", style=styles["body_no_indent"])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                remember(paragraph)
            else:
                paragraph = document.add_paragraph(style=styles["body_no_indent"])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Cm(14.5))
                remember(paragraph)
                caption_text = image_match.group("title") or image_match.group("alt")
                if caption_text:
                    figure_counters[chapter] = figure_counters.get(chapter, 0) + 1
                    caption = document.add_paragraph(
                        f"图 {chapter}-{figure_counters[chapter]} {caption_text}",
                        style=styles["caption"],
                    )
                    remember(caption)
            index += 1
            continue

        # Markdown table: current row + separator + data rows.
        if stripped.startswith("|") and index + 1 < len(lines) and TABLE_DIVIDER_RE.match(lines[index + 1]):
            flush_paragraph()
            header = _split_table_row(line)
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_split_table_row(lines[index]))
                index += 1
            table = document.add_table(rows=1, cols=len(header))
            for col, value in enumerate(header):
                table.rows[0].cells[col].text = value
            for row_values in rows:
                row = table.add_row()
                normalized = row_values + [""] * max(0, len(header) - len(row_values))
                for col, value in enumerate(normalized[: len(header)]):
                    row.cells[col].text = value
            _format_table(table, styles["table"])
            # Tables are block elements; remember their nearest preceding/following
            # paragraph for bookmark boundaries. Add a zero-height paragraph after.
            anchor = document.add_paragraph(style=styles["body_no_indent"])
            anchor.paragraph_format.space_before = Pt(0)
            anchor.paragraph_format.space_after = Pt(0)
            anchor.paragraph_format.line_spacing = 1.0
            remember(anchor)
            continue

        list_match = LIST_RE.match(line)
        if list_match:
            flush_paragraph()
            paragraph = document.add_paragraph(_clean_inline(list_match.group("text")), style=styles["body_no_indent"])
            indent_level = len(list_match.group("indent")) // 2
            paragraph.paragraph_format.left_indent = Pt(24 + indent_level * 18)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            marker = "•" if not list_match.group("mark")[0].isdigit() else list_match.group("mark").strip()
            paragraph.text = f"{marker} {_clean_inline(list_match.group('text'))}"
            remember(paragraph)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
        else:
            paragraph_buffer.append(line)
        index += 1

    flush_paragraph()
    if in_code:
        paragraph = document.add_paragraph("\n".join(code_lines), style=styles["code"])
        remember(paragraph)

    if first_body_paragraph is None or last_body_paragraph is None:
        raise RenderError("Markdown 正文为空")
    return first_body_paragraph, last_body_paragraph


def _front_payload(markdown: MarkdownDocument) -> dict:
    metadata = markdown.metadata
    document = metadata.get("document", {}) if isinstance(metadata.get("document"), dict) else {}
    software = metadata.get("software", {}) if isinstance(metadata.get("software"), dict) else {}
    front = metadata.get("front_matter", {}) if isinstance(metadata.get("front_matter"), dict) else {}
    return {
        "archive_id": front.get("archive_id", ""),
        "classification": metadata.get("classification", ""),
        "project_code": front.get("project_code", ""),
        "document_id": document.get("id", ""),
        "phase": front.get("phase", ""),
        "document_version": document.get("version", ""),
        "software_name": software.get("name", ""),
        "document_title": document.get("title", ""),
        "organization": metadata.get("organization", ""),
        "date_cn": front.get("date_cn", metadata.get("date", "")),
        "signatures": metadata.get("signatures", {}),
        "revisions": metadata.get("revisions", []),
    }


def _normalized_bookmark_text(document_xml: bytes) -> str:
    root = etree.fromstring(document_xml)
    body = root.find("./w:body", NS)
    if body is None:
        return ""

    starts = root.xpath(
        './/w:bookmarkStart[@w:name=$name]',
        namespaces=NS,
        name=BOOKMARK_NAME,
    )
    if len(starts) != 1:
        return ""
    body_bookmark_id = starts[0].get(f"{{{W}}}id")
    if body_bookmark_id is None:
        return ""

    # The body range may legally contain nested bookmarks. TOC finalization adds
    # one bookmark to every heading so its cached hyperlinks have real targets.
    # Stop only at the bookmarkEnd whose id matches GJB_BODY; stopping at the
    # first nested bookmarkEnd truncates the normalized body text and makes an
    # otherwise content-preserving TOC refresh look like a Word edit.
    active = False
    pieces: list[str] = []
    for element in body.iter():
        if not active:
            if (
                element.tag == f"{{{W}}}bookmarkStart"
                and element.get(f"{{{W}}}name") == BOOKMARK_NAME
                and element.get(f"{{{W}}}id") == body_bookmark_id
            ):
                active = True
            continue
        if (
            element.tag == f"{{{W}}}bookmarkEnd"
            and element.get(f"{{{W}}}id") == body_bookmark_id
        ):
            break
        if element.tag == f"{{{W}}}t" and element.text:
            pieces.append(element.text)
    return re.sub(r"\s+", " ", " ".join(pieces)).strip()


def _patch_settings_with_source(docx_path: Path, markdown_source: str) -> None:
    with tempfile.TemporaryDirectory(prefix="gjb438c-docvars-") as temp_name:
        temp = Path(temp_name)
        with ZipFile(docx_path) as archive:
            archive.extractall(temp)
        settings_path = temp / "word" / "settings.xml"
        settings_tree = etree.parse(str(settings_path))
        settings_root = settings_tree.getroot()
        update = settings_root.find("./w:updateFields", NS)
        if update is None:
            update = etree.SubElement(settings_root, f"{{{W}}}updateFields")
        update.set(f"{{{W}}}val", "true")

        doc_vars = settings_root.find("./w:docVars", NS)
        if doc_vars is None:
            doc_vars = etree.SubElement(settings_root, f"{{{W}}}docVars")
        for variable in list(doc_vars):
            name = variable.get(f"{{{W}}}name", "")
            if name.startswith(DOCVAR_PREFIX) or name in {DOCVAR_HASH, DOCVAR_SOURCE_HASH}:
                doc_vars.remove(variable)

        compressed = gzip.compress(markdown_source.encode("utf-8"), compresslevel=9)
        encoded = base64.b64encode(compressed).decode("ascii")
        chunks = [encoded[index : index + 28000] for index in range(0, len(encoded), 28000)]
        for index, chunk in enumerate(chunks):
            variable = etree.SubElement(doc_vars, f"{{{W}}}docVar")
            variable.set(f"{{{W}}}name", f"{DOCVAR_PREFIX}{index:04d}")
            variable.set(f"{{{W}}}val", chunk)
        document_xml = (temp / "word" / "document.xml").read_bytes()
        body_text = _normalized_bookmark_text(document_xml)
        for name, value in (
            (DOCVAR_HASH, sha256(body_text.encode("utf-8")).hexdigest()),
            (DOCVAR_SOURCE_HASH, sha256(markdown_source.encode("utf-8")).hexdigest()),
        ):
            variable = etree.SubElement(doc_vars, f"{{{W}}}docVar")
            variable.set(f"{{{W}}}name", name)
            variable.set(f"{{{W}}}val", value)

        settings_tree.write(
            str(settings_path), xml_declaration=True, encoding="UTF-8", standalone="yes"
        )
        rebuilt = docx_path.with_suffix(docx_path.suffix + ".tmp")
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as archive:
            for file in temp.rglob("*"):
                if file.is_file():
                    archive.write(file, file.relative_to(temp).as_posix())
        rebuilt.replace(docx_path)


def render_document(
    markdown_path: str | Path,
    output_path: str | Path,
    *,
    profile: str = "review",
    baseline_srs: str | Path | None = None,
    front_template: str | Path | None = None,
) -> RenderResult:
    source = Path(markdown_path)
    output = Path(output_path)
    markdown = parse_markdown(source)
    report = audit_markdown(source, profile=profile, baseline_srs=baseline_srs)
    if not report.passed:
        raise RenderError(report.to_text())
    release = profile == "release"

    template = Path(front_template) if front_template else default_front_matter_template()
    configured = nested_get(markdown.metadata, "front_matter.template")
    if front_template is None and configured:
        candidate = Path(str(configured))
        if not candidate.is_absolute():
            for base in (source.parent, Path(__file__).resolve().parents[1]):
                resolved = (base / candidate).resolve()
                if resolved.is_file():
                    candidate = resolved
                    break
        if candidate.is_file():
            template = candidate

    with tempfile.TemporaryDirectory(prefix="gjb438c-render-") as temp_name:
        temp = Path(temp_name)
        front_docx = temp / "front.docx"
        try:
            render_front_matter(template, _front_payload(markdown), front_docx, release=release)
        except FrontMatterError as exc:
            raise RenderError(str(exc)) from exc

        document = Document(front_docx)
        styles = configure_styles(document)
        # Preserve the official first three pages; Python-docx adds the section
        # break after the existing revision table without rewriting those pages.
        toc_section = document.add_section(WD_SECTION.NEW_PAGE)
        _set_section_a4(toc_section)
        _set_page_numbering(toc_section, fmt="upperRoman", start=1)
        _set_footer_page(toc_section, roman=True)

        title = document.add_paragraph("目录", style=styles["toc_title"])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_paragraph = document.add_paragraph(style=styles["body_no_indent"])
        _add_toc(toc_paragraph)

        body_section = document.add_section(WD_SECTION.NEW_PAGE)
        _set_section_a4(body_section)
        _set_page_numbering(body_section, fmt="decimal", start=1)
        _set_footer_page(body_section, roman=False)

        first, last = _render_markdown_body(
            document, markdown, styles, release=release
        )
        _bookmark_start(first)
        _bookmark_end(last)

        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        _patch_settings_with_source(output, markdown.raw)
    return RenderResult(output, report)
